#!/usr/bin/env python3
import json
import math
import os
import re
import sys
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple


def mm_to_pt(value_mm: float) -> float:
    return float(value_mm) * 72.0 / 25.4


def parse_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    if isinstance(value, (int, float)):
        return bool(value)
    text = str(value).strip().lower()
    if text in ("1", "true", "yes", "y", "on"):
        return True
    if text in ("0", "false", "no", "n", "off", ""):
        return False
    return False


def fmt_amount(amount: Any, *, decimal_comma: bool) -> str:
    text = f"{float(amount or 0):.2f}"
    return text.replace(".", ",") if decimal_comma else text


def safe_filename_part(value: Any) -> str:
    text = str(value or "").strip()
    text = text.replace("\\", "-").replace("/", "-")
    text = re.sub(r"[^0-9A-Za-z._-]+", "_", text)
    text = text.strip("._-")
    return text or "unknown"


def parse_input_json() -> Dict[str, Any]:
    raw = sys.stdin.buffer.read()
    if not raw:
        raise ValueError("No JSON input received on stdin")
    try:
        return json.loads(raw.decode("utf-8"))
    except Exception as exc:
        raise ValueError(f"Invalid JSON input: {exc}") from exc


def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def number_to_words_fr(num: float) -> str:
    whole_part = int(math.floor(float(num or 0)))
    if whole_part == 0:
        return "Zero"

    units = ["", "Un", "Deux", "Trois", "Quatre", "Cinq", "Six", "Sept", "Huit", "Neuf"]
    teens = [
        "Dix",
        "Onze",
        "Douze",
        "Treize",
        "Quatorze",
        "Quinze",
        "Seize",
        "Dix Sept",
        "Dix Huit",
        "Dix Neuf",
    ]
    tens = [
        "",
        "",
        "Vingt",
        "Trente",
        "Quarante",
        "Cinquante",
        "Soixante",
        "Soixante Dix",
        "Quatre Vingt",
        "Quatre Vingt Dix",
    ]

    def convert_below_hundred(n: int) -> str:
        if n < 10:
            return units[n]
        if n < 20:
            return teens[n - 10]
        tens_digit = n // 10
        unit_digit = n % 10
        tens_word = tens[tens_digit]
        if unit_digit == 0:
            return tens_word
        if unit_digit == 1 and tens_digit in (2, 3, 4, 5, 6):
            return f"{tens_word} Et {units[unit_digit]}".strip()
        unit_word = units[unit_digit]
        return f"{tens_word} {unit_word}".strip()

    def convert_below_thousand(n: int) -> str:
        hundreds = n // 100
        remainder = n % 100
        words = ""
        if hundreds > 0:
            words += f"{units[hundreds]} Cent".strip()
            if remainder > 0:
                words += " "
        if remainder > 0:
            words += convert_below_hundred(remainder)
        return words.strip()

    def convert(n: int) -> str:
        if n < 1000:
            return convert_below_thousand(n)
        if n < 1_000_000:
            thousands = n // 1000
            remainder = n % 1000
            thousand_words = "Mille" if thousands == 1 else f"{convert_below_thousand(thousands)} Mille"
            remainder_words = f" {convert_below_thousand(remainder)}" if remainder > 0 else ""
            return f"{thousand_words}{remainder_words}".strip()
        millions = n // 1_000_000
        remainder = n % 1_000_000
        million_words = "Un Million" if millions == 1 else f"{convert_below_thousand(millions)} Millions"
        remainder_words = f" {convert(remainder)}" if remainder > 0 else ""
        return f"{million_words}{remainder_words}".strip()

    return convert(whole_part)


@dataclass(frozen=True)
class Layout:
    page_margin_mm: float = 12.0
    page_top_margin_mm: float = 10.0
    page_bottom_margin_mm: float = 10.0
    page_left_margin_mm: float = 12.0
    page_right_margin_mm: float = 12.0
    header_col_gap_mm: float = 6.0
    table_top_gap_mm: float = 6.0
    table_header_h_mm: float = 17.0
    table_row_h_mm: float = 13.0
    table_rows_per_page: int = 8
    table_rows_page1: int = 8
    table_rows_continuation: int = 6
    post_table_top_gap_mm: float = 12.0
    post_grid_gap_mm: float = 10.0
    post_indent_mm: float = 6.0
    post_block_margin_top_mm: float = 6.0
    post_bottom_padding_mm: float = 12.0
    post_signature_reserve_mm: float = 32.0


def build_role_table(
    *,
    rows: List[List[Any]],
    col_widths_mm: List[float],
    row_heights_mm: List[float],
    include_totals_row: bool,
):
    raise RuntimeError("PDF generation removed (DOCX-only).")

    header_style = ParagraphStyle(
        "header",
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=10.5,
        alignment=1,  # center
    )
    body_style_center = ParagraphStyle(
        "body_center",
        fontName="Helvetica",
        fontSize=10,
        leading=12.5,
        alignment=1,
    )
    body_style_left = ParagraphStyle(
        "body_left",
        fontName="Helvetica",
        fontSize=10,
        leading=12.5,
        alignment=0,
    )
    cin_style = ParagraphStyle(
        "cin_left",
        fontName="Helvetica",
        fontSize=7,
        leading=8.5,
        alignment=0,
    )

    header_row = [
        "N° DESP. D'ATTACH.",
        Paragraph("PRENOMS ET NOMS", header_style),
        Paragraph("EMPLOIS", header_style),
        Paragraph("NOMBRE DE<br/>JOURNEES", header_style),
        Paragraph("PRIX DE LA<br/>JOURNEE", header_style),
        Paragraph("BRUT A PAYER", header_style),
        Paragraph("PRELEVEMENT<br/>I.G.R", header_style),
        Paragraph("NET A PAYER", header_style),
        Paragraph("N° DE LA C.I.N<br/>ET SIGNATURE", header_style),
    ]

    body_rows: List[List[Any]] = []
    for row in rows:
        body_rows.append(
            [
                Paragraph(str(row[0] or ""), body_style_center),
                Paragraph(str(row[1] or ""), body_style_left),
                Paragraph(str(row[2] or ""), body_style_center),
                Paragraph(str(row[3] or ""), body_style_center),
                Paragraph(str(row[4] or ""), body_style_center),
                Paragraph(str(row[5] or ""), body_style_center),
                Paragraph(str(row[6] or ""), body_style_center),
                Paragraph(str(row[7] or ""), body_style_center),
                Paragraph(str(row[8] or ""), cin_style),
            ]
        )

    data = [header_row] + body_rows

    col_widths_pt = [mm_to_pt(w) for w in col_widths_mm]
    row_heights_pt = [mm_to_pt(h) for h in row_heights_mm]

    table = Table(data, colWidths=col_widths_pt, rowHeights=row_heights_pt, repeatRows=1)
    style_cmds = [
        ("GRID", (0, 0), (-1, -1), 0.8, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, 0), mm_to_pt(2.0)),
        ("RIGHTPADDING", (0, 0), (-1, 0), mm_to_pt(2.0)),
        ("TOPPADDING", (0, 0), (-1, 0), mm_to_pt(2.0)),
        ("BOTTOMPADDING", (0, 0), (-1, 0), mm_to_pt(2.0)),
        ("LEFTPADDING", (0, 1), (-1, -1), mm_to_pt(3.0)),
        ("RIGHTPADDING", (0, 1), (-1, -1), mm_to_pt(3.0)),
        ("TOPPADDING", (0, 1), (-1, -1), mm_to_pt(2.0)),
        ("BOTTOMPADDING", (0, 1), (-1, -1), mm_to_pt(2.0)),
        ("ALIGN", (1, 1), (1, -1), "LEFT"),
        ("ALIGN", (8, 1), (8, -1), "LEFT"),
        ("TEXTANGLE", (0, 0), (0, 0), 90),
        ("FONTSIZE", (0, 0), (0, 0), 7),
    ]

    if include_totals_row:
        style_cmds.extend(
            [
                ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, -1), (-1, -1), 10),
            ]
        )

    table.setStyle(TableStyle(style_cmds))
    return table


def draw_underlined_text(c, x: float, y: float, text: str, font: str, size: float) -> float:
    c.setFont(font, size)
    c.drawString(x, y, text)
    width = c.stringWidth(text, font, size)
    c.setLineWidth(0.8)
    c.line(x, y - 1.2, x + width, y - 1.2)
    return width


def draw_role_pdf(payload: Dict[str, Any], pdf_path: str) -> None:
    raise RuntimeError("PDF generation removed (DOCX-only).")

    layout = Layout()
    page_w, page_h = A4
    margin = mm_to_pt(layout.page_margin_mm)
    content_x = margin
    content_w = page_w - 2 * margin

    # Column widths normalized to match the HTML colgroup ratios (sum=101% there)
    col_perc = [6, 24, 7, 8, 8, 9, 9, 9, 21]
    perc_total = float(sum(col_perc))
    col_widths_mm = [(content_w / mm_to_pt(1.0)) * (p / perc_total) for p in col_perc]

    c = canvas.Canvas(pdf_path, pagesize=A4)
    regisseur = str(payload.get("regisseurName") or "MAJDA TAKNOUTI")
    year = str(payload.get("year") or "")
    reference = payload.get("referenceValues") or {}
    decimal_comma = parse_bool(payload.get("decimalComma"))

    sections = payload.get("sections") or []
    page1_rows_per_page = int(payload.get("splitIndex") or layout.table_rows_page1 or layout.table_rows_per_page)
    continuation_rows_per_page = int(payload.get("continuationRows") or layout.table_rows_continuation or page1_rows_per_page)

    header_left_lines = [
        "ROYAUME DU MAROC",
        "MINISTERE DE L'INTERIEUR",
        "REGIE DE DEPENSES AUPRES DE",
        "LA COMMUNE OULED NACEUR",
        "…………………………..",
    ]
    header_center_lines = ["DEPENSES EN REGIE", "SALAIRE DU PERSONNEL OCCASIONNEL"]
    header_right = "ANNEXE : 9………."

    for section_index, section in enumerate(sections):
        section_workers = section.get("workers") or []
        start_date = str(section.get("startDate") or "")
        end_date = str(section.get("endDate") or "")

        # Split workers for page 1 and continuation pages.
        page1_workers = section_workers[:page1_rows_per_page]
        remaining = section_workers[page1_rows_per_page:]
        remaining_chunks = [
            remaining[i : i + continuation_rows_per_page]
            for i in range(0, len(remaining), continuation_rows_per_page)
        ]
        if not remaining_chunks:
            remaining_chunks = [[]]  # always produce the Page 2 template

        # --- PAGE 1 ---
        y = page_h - margin

        # Header (3 columns)
        col_gap = mm_to_pt(layout.header_col_gap_mm)
        col_w = (content_w - 2 * col_gap) / 3.0
        left_x = content_x
        center_x = content_x + col_w + col_gap
        right_x = content_x + 2 * (col_w + col_gap)

        header_font = "Helvetica-Bold"
        header_size = 10
        leading = header_size * 1.15

        y0 = y - header_size
        for i, line in enumerate(header_left_lines):
            c.setFont(header_font, header_size)
            c.drawString(left_x, y0 - i * leading, line)

        for i, line in enumerate(header_center_lines):
            c.setFont(header_font, header_size)
            c.drawCentredString(center_x + col_w / 2.0, y0 - i * leading, line)

        c.setFont("Helvetica-Bold", 10)
        c.drawRightString(right_x + col_w, y0, header_right)

        header_height = max(len(header_left_lines), len(header_center_lines), 1) * leading + header_size * 0.25
        y = y - header_height - mm_to_pt(3.0)

        # Title
        title = "ROLE DES JOURNEES D'OUVRIERS EMPLOYES"
        c.setFont("Helvetica-Bold", 11)
        title_y = y - 11
        c.drawCentredString(content_x + content_w / 2.0, title_y, title)
        title_w = c.stringWidth(title, "Helvetica-Bold", 11)
        c.setLineWidth(0.8)
        c.line((content_x + content_w / 2.0) - title_w / 2.0, title_y - 1.2, (content_x + content_w / 2.0) + title_w / 2.0, title_y - 1.2)
        y = title_y - mm_to_pt(4.0)

        # Info section
        info_size = 10
        info_leading = info_size * 1.25
        c.setFont("Helvetica", info_size)

        # Row 1: Regisseur + DU/AU boxes
        y = y - info_size
        label = "NOM DU REGISSEUR :"
        label_w = draw_underlined_text(c, content_x, y, label, "Helvetica-Bold", info_size)
        c.setFont("Helvetica", info_size)
        c.drawString(content_x + label_w + mm_to_pt(2.0), y, regisseur)

        box_w = mm_to_pt(25.0)
        box_h = mm_to_pt(6.0)
        label_gap = mm_to_pt(2.0)
        group_gap = mm_to_pt(6.0)

        au_label = "AU :"
        du_label = "DU :"
        c.setFont("Helvetica-Bold", info_size)
        au_label_w = c.stringWidth(au_label, "Helvetica-Bold", info_size)
        du_label_w = c.stringWidth(du_label, "Helvetica-Bold", info_size)

        au_group_w = au_label_w + label_gap + box_w
        au_x = content_x + content_w - au_group_w
        du_group_w = du_label_w + label_gap + box_w
        du_x = au_x - group_gap - du_group_w

        # DU
        c.drawString(du_x, y, du_label)
        c.rect(du_x + du_label_w + label_gap, y - mm_to_pt(1.5), box_w, box_h, stroke=1, fill=0)
        c.setFont("Helvetica", info_size)
        c.drawCentredString(du_x + du_label_w + label_gap + box_w / 2.0, y, start_date)

        # AU
        c.setFont("Helvetica-Bold", info_size)
        c.drawString(au_x, y, au_label)
        c.rect(au_x + au_label_w + label_gap, y - mm_to_pt(1.5), box_w, box_h, stroke=1, fill=0)
        c.setFont("Helvetica", info_size)
        c.drawCentredString(au_x + au_label_w + label_gap + box_w / 2.0, y, end_date)

        y -= info_leading

        # Row 2
        c.setFont("Helvetica-Bold", info_size)
        c.drawString(content_x, y, "(2)TRAVAUX DIVERS A LA COMMUNE OULED NACEUR")
        y -= info_leading

        # Row 3 (reference values)
        items = [
            ("ANNEE :", year),
            ("CHAP :", str(reference.get("chapitre") or reference.get("chap") or "")),
            ("Art :", str(reference.get("article") or reference.get("art") or "")),
            ("Prog :", str(reference.get("programme") or reference.get("prog") or "")),
            ("Proj :", str(reference.get("projet") or reference.get("proj") or "")),
            ("Ligne :", str(reference.get("ligne") or "")),
        ]
        x_cursor = content_x
        for label_text, value_text in items:
            c.setFont("Helvetica-Bold", info_size)
            lw = c.stringWidth(label_text, "Helvetica-Bold", info_size)
            vw = c.stringWidth(value_text, "Helvetica", info_size)
            segment_w = lw + vw + mm_to_pt(6.0)
            if x_cursor + segment_w > content_x + content_w:
                x_cursor = content_x
                y -= info_leading
            c.drawString(x_cursor, y, label_text)
            c.setFont("Helvetica", info_size)
            c.drawString(x_cursor + lw, y, value_text)
            x_cursor += segment_w
        y -= info_leading

        # Row 4
        somme_label = "SOMME A PAYER :"
        somme_value = fmt_amount(section.get("totalNet") or 0, decimal_comma=decimal_comma)
        c.setFont("Helvetica-Bold", info_size)
        slw = c.stringWidth(somme_label, "Helvetica-Bold", info_size)
        c.drawString(content_x, y, somme_label)
        c.setFont("Helvetica", info_size)
        c.drawString(content_x + slw + mm_to_pt(2.0), y, somme_value)

        # Table gap
        y -= mm_to_pt(layout.table_top_gap_mm) + info_leading

        def build_rows_for_workers(
            workers_slice: List[Dict[str, Any]],
            start_row_number: int,
            pad_to_rows: Optional[int] = None,
        ) -> List[List[Any]]:
            rows_out = []
            for idx, w in enumerate(workers_slice):
                type_display = "O.S" if str(w.get("type") or "").strip() == "OS" else "O.N.S"
                cin = str(w.get("cin") or "")
                cin_validite = str(w.get("cin_validite") or "")
                cin_block = f"CIN N°: {cin}<br/>AU: {cin_validite}"
                rows_out.append(
                    [
                        str(start_row_number + idx),
                        str(w.get("nom_prenom") or ""),
                        type_display,
                        str(int(w.get("daysWorked") or 0)),
                        fmt_amount(w.get("salaire_journalier") or 0, decimal_comma=decimal_comma),
                        fmt_amount(w.get("grossSalary") or 0, decimal_comma=decimal_comma),
                        fmt_amount(w.get("deduction") or 0, decimal_comma=decimal_comma),
                        fmt_amount(w.get("netSalary") or 0, decimal_comma=decimal_comma),
                        cin_block,
                    ]
                )

            if pad_to_rows is not None:
                while len(rows_out) < pad_to_rows:
                    row_number = start_row_number + len(rows_out)
                    rows_out.append([str(row_number), "", "", "", "", "", "", "", ""])
            return rows_out

        page1_rows = build_rows_for_workers(page1_workers, 1, pad_to_rows=page1_rows_per_page)
        table_row_heights_mm = [layout.table_header_h_mm] + [layout.table_row_h_mm] * max(len(page1_rows), 1)
        table = build_role_table(
            rows=page1_rows or [["", "", "", "", "", "", "", "", ""]],
            col_widths_mm=col_widths_mm,
            row_heights_mm=table_row_heights_mm,
            include_totals_row=False,
        )
        tw, th = table.wrap(0, 0)
        table.drawOn(c, content_x, y - th)

        c.showPage()

        # --- CONTINUATION PAGES (PAGE 2+) ---
        for chunk_index, chunk in enumerate(remaining_chunks):
            is_last_chunk = chunk_index == (len(remaining_chunks) - 1)
            y2 = page_h - margin
            y2 -= mm_to_pt(layout.table_top_gap_mm)

            start_row_number = page1_rows_per_page + 1 + chunk_index * continuation_rows_per_page
            body_rows = build_rows_for_workers(chunk, start_row_number, pad_to_rows=continuation_rows_per_page)

            include_totals = bool(is_last_chunk)
            if include_totals:
                body_rows.append(
                    [
                        "",
                        "TOTAL:",
                        "",
                        str(int(section.get("totalDays") or 0)),
                        "",
                        fmt_amount(section.get("totalGross") or 0, decimal_comma=decimal_comma),
                        fmt_amount(section.get("totalDeduction") or 0, decimal_comma=decimal_comma),
                        fmt_amount(section.get("totalNet") or 0, decimal_comma=decimal_comma),
                        "",
                    ]
                )

            row_heights_mm = [layout.table_header_h_mm] + [layout.table_row_h_mm] * (len(body_rows))
            table2 = build_role_table(
                rows=body_rows,
                col_widths_mm=col_widths_mm,
                row_heights_mm=row_heights_mm,
                include_totals_row=include_totals,
            )
            tw2, th2 = table2.wrap(0, 0)
            table2.drawOn(c, content_x, y2 - th2)
            y_after_table = y2 - th2

            if include_totals:
                # Post-table block
                post_y = y_after_table - mm_to_pt(layout.post_table_top_gap_mm)
                post_font_size = 8
                post_leading = post_font_size * 1.35

                gap = mm_to_pt(layout.post_grid_gap_mm)
                left_w = (content_w - gap) * 0.58
                right_w = (content_w - gap) * 0.42
                left_x = content_x
                right_x = content_x + left_w + gap

                pay_date = str(section.get("payDate") or section.get("documentDate") or end_date)
                total_net = float(section.get("totalNet") or 0)
                words = number_to_words_fr(total_net).upper()
                cents = int(round((total_net - math.floor(total_net)) * 100.0))
                cents_str = str(cents).rjust(2, "0")

                # Left column
                c.setFont("Helvetica-Bold", post_font_size)
                c.drawString(left_x, post_y, "NOUS SOUSSIGNONS :")

                post_y -= mm_to_pt(layout.post_indent_mm)  # vertical spacing similar to template
                c.setFont("Helvetica-Bold", post_font_size)
                c.drawString(left_x + mm_to_pt(layout.post_indent_mm), post_y, "Mr")

                post_y -= mm_to_pt(layout.post_indent_mm)
                c.drawString(left_x + mm_to_pt(layout.post_indent_mm), post_y, "Mr")

                post_y -= post_leading
                c.setFont("Helvetica-Bold", post_font_size)
                c.drawString(left_x, post_y, "CERTIFIONS QUE LES SIEURS :")

                post_y -= post_leading
                c.setFont("Helvetica", post_font_size)
                c.drawCentredString(left_x + left_w / 2.0, post_y, "Portés au présent Role ont été payés en notre")
                post_y -= post_leading
                c.drawCentredString(left_x + left_w / 2.0, post_y, "présence, après apposition de leurs signatures")

                post_y -= mm_to_pt(3.0) + post_leading
                c.setFont("Helvetica-Bold", post_font_size)
                c.drawString(left_x, post_y, f"PAYER PAR Moi Le : {pay_date}")

                # Right column
                right_y = y_after_table - mm_to_pt(layout.post_table_top_gap_mm)
                c.setFont("Helvetica", post_font_size)
                c.drawString(right_x, right_y, "LE PRESENT ROLE S'ELEVANT A LA SOMME DE :")
                right_y -= post_leading
                c.setFont("Helvetica-Bold", post_font_size)
                c.drawString(right_x, right_y, f"{words} DHS {cents_str} CTS.")

                right_y -= mm_to_pt(10.0) + post_leading
                president = "LE PRESIDENT DU CONSEIL"
                c.setFont("Helvetica-Bold", post_font_size)
                pres_y = right_y
                c.drawCentredString(right_x + right_w / 2.0, pres_y, president)
                pres_w = c.stringWidth(president, "Helvetica-Bold", post_font_size)
                c.setLineWidth(0.8)
                c.line((right_x + right_w / 2.0) - pres_w / 2.0, pres_y - 1.2, (right_x + right_w / 2.0) + pres_w / 2.0, pres_y - 1.2)

                # Regisseur line at bottom of block
                reg_line_y = pres_y - mm_to_pt(10.0)
                reg_text = "LE REGISSEUR DE DEPENSES"
                c.setFont("Helvetica-Bold", post_font_size)
                c.drawString(content_x, reg_line_y, reg_text)
                reg_w = c.stringWidth(reg_text, "Helvetica-Bold", post_font_size)
                c.setLineWidth(0.8)
                c.line(content_x, reg_line_y - 1.2, content_x + reg_w, reg_line_y - 1.2)

            if not (is_last_chunk and section_index == (len(sections) - 1)):
                c.showPage()

    c.save()


def set_docx_cell_border(cell, **kwargs) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = qn(f"w:{edge}")
            element = tc_borders.find(tag)
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tc_borders.append(element)
            element.set(qn("w:val"), edge_data.get("val", "single"))
            element.set(qn("w:sz"), str(edge_data.get("sz", 8)))
            element.set(qn("w:color"), edge_data.get("color", "000000"))


def set_docx_cell_text_direction(cell, direction: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    text_dir = tc_pr.find(qn("w:textDirection"))
    if text_dir is None:
        text_dir = OxmlElement("w:textDirection")
        tc_pr.append(text_dir)
    text_dir.set(qn("w:val"), direction)


def remove_docx_table_borders(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "nil")


def mm_to_twips(value_mm: float) -> int:
    # 1 point = 20 twips
    return int(round(mm_to_pt(float(value_mm)) * 20.0))


def set_page_margins(section, *, top_mm: float, bottom_mm: float, left_mm: float, right_mm: float) -> None:
    from docx.shared import Mm

    section.top_margin = Mm(top_mm)
    section.bottom_margin = Mm(bottom_mm)
    section.left_margin = Mm(left_mm)
    section.right_margin = Mm(right_mm)


def set_default_font(doc, *, font_name: str, font_size_pt: float) -> None:
    from docx.shared import Pt

    normal_style = doc.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size_pt)


def add_centered_title(doc, text: str, *, font_size_pt: float) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(font_size_pt)


def add_table_with_widths(doc, *, cols: int, col_widths_mm: List[float]):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Mm

    t = doc.add_table(rows=1, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    for i, w in enumerate(col_widths_mm):
        if i < len(t.columns):
            t.columns[i].width = Mm(w)
    return t


def set_docx_table_cell_margins(table, *, top_mm: float, bottom_mm: float, left_mm: float, right_mm: float) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)

    for edge, mm_val in (("top", top_mm), ("bottom", bottom_mm), ("left", left_mm), ("right", right_mm)):
        node = cell_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            cell_mar.append(node)
        node.set(qn("w:w"), str(mm_to_twips(mm_val)))
        node.set(qn("w:type"), "dxa")


def set_docx_table_fixed_layout(table, *, total_width_mm: float, col_widths_mm: List[float]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(mm_to_twips(total_width_mm)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")

    for col_idx, width_mm in enumerate(col_widths_mm):
        if col_idx >= len(table.columns):
            break
        table.columns[col_idx].width = Mm(width_mm)
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = Mm(width_mm)


def format_doc_date(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    if re.match(r"^\d{2}/\d{2}/\d{4}$", text):
        return text
    iso_match = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", text)
    if iso_match:
        return f"{iso_match.group(3)}/{iso_match.group(2)}/{iso_match.group(1)}"
    return text


def draw_role_docx(payload: Dict[str, Any], docx_path: str) -> None:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    layout = Layout()

    doc = Document()
    section0 = doc.sections[0]
    section0.page_width = Mm(210)
    section0.page_height = Mm(297)
    set_page_margins(
        section0,
        top_mm=layout.page_top_margin_mm,
        bottom_mm=layout.page_bottom_margin_mm,
        left_mm=layout.page_left_margin_mm,
        right_mm=layout.page_right_margin_mm,
    )
    set_default_font(doc, font_name="Times New Roman", font_size_pt=9)

    regisseur = str(payload.get("regisseurName") or "MAJDA TAKNOUTI")
    year = str(payload.get("year") or "")
    reference = payload.get("referenceValues") or {}
    sections = payload.get("sections") or []
    if not sections:
        raise ValueError("Missing required data: sections")

    decimal_comma = parse_bool(payload.get("decimalComma"))
    page1_worker_rows = max(1, int(payload.get("splitIndex") or layout.table_rows_page1 or layout.table_rows_per_page))
    continuation_worker_rows = max(1, int(payload.get("continuationRows") or layout.table_rows_continuation or layout.table_rows_page1))

    usable_w_mm = 210 - layout.page_left_margin_mm - layout.page_right_margin_mm
    gap_mm = layout.post_grid_gap_mm
    left_block_mm = (usable_w_mm - gap_mm) * 0.56
    right_block_mm = (usable_w_mm - gap_mm) * 0.44

    # Column ratios tuned to match the scanned paper with full-width fixed table layout.
    col_perc = [6, 27, 9, 8, 8, 10, 10, 10, 12]
    perc_total = float(sum(col_perc))
    col_widths_mm = [round(usable_w_mm * (p / perc_total), 3) for p in col_perc]
    if col_widths_mm:
        col_widths_mm[-1] = round(usable_w_mm - sum(col_widths_mm[:-1]), 3)

    def to_float(value: Any, default: float = 0.0) -> float:
        try:
            return float(value)
        except Exception:
            return float(default)

    def round2(value: Any) -> float:
        return float(f"{to_float(value, 0):.2f}")

    def format_cell_text(cell, text: str, *, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size_pt=9, underline=False) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = bold
        run.underline = underline
        run.font.size = Pt(size_pt)

    def sum_workers(rows: List[Dict[str, Any]]) -> Dict[str, float]:
        totals_out = {"totalDays": 0.0, "totalGross": 0.0, "totalDeduction": 0.0, "totalNet": 0.0}
        for row in rows:
            totals_out["totalDays"] += to_float(row.get("daysWorked"), 0)
            totals_out["totalGross"] += to_float(row.get("grossSalary"), 0)
            totals_out["totalDeduction"] += to_float(row.get("deduction"), 0)
            totals_out["totalNet"] += to_float(row.get("netSalary"), 0)
        totals_out["totalGross"] = round2(totals_out["totalGross"])
        totals_out["totalDeduction"] = round2(totals_out["totalDeduction"])
        totals_out["totalNet"] = round2(totals_out["totalNet"])
        return totals_out

    def normalize_totals(raw: Dict[str, Any], fallback: Dict[str, float]) -> Dict[str, float]:
        total_days = to_float(raw.get("totalDays"), fallback["totalDays"])
        total_gross = round2(raw.get("totalGross") if raw.get("totalGross") is not None else fallback["totalGross"])
        total_deduction = round2(raw.get("totalDeduction") if raw.get("totalDeduction") is not None else fallback["totalDeduction"])
        total_net = round2(raw.get("totalNet") if raw.get("totalNet") is not None else fallback["totalNet"])
        return {
            "totalDays": total_days,
            "totalGross": total_gross,
            "totalDeduction": total_deduction,
            "totalNet": total_net,
        }

    def add_role_header_page1(doc_ref, data: Dict[str, Any]) -> None:
        header_tbl = doc_ref.add_table(rows=1, cols=3)
        header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_docx_table_borders(header_tbl)
        set_docx_table_fixed_layout(
            header_tbl,
            total_width_mm=usable_w_mm,
            col_widths_mm=[usable_w_mm / 3.0, usable_w_mm / 3.0, usable_w_mm / 3.0],
        )

        left_cell, center_cell, right_cell = header_tbl.rows[0].cells
        for c in (left_cell, center_cell, right_cell):
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        left_p = left_cell.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_p.paragraph_format.space_before = Pt(0)
        left_p.paragraph_format.space_after = Pt(0)
        for line in [
            "ROYAUME DU MAROC",
            "MINISTERE DE L'INTERIEUR",
            "REGIE DE DEPENSES AUPRES DE",
            "LA COMMUNE OULED NACEUR",
            "…………………………..",
        ]:
            r = left_p.add_run(line + "\n")
            r.bold = True
            r.font.size = Pt(9)

        center_p = center_cell.paragraphs[0]
        center_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        center_p.paragraph_format.space_before = Pt(0)
        center_p.paragraph_format.space_after = Pt(0)
        for line in ["DEPENSES EN REGIE", "SALAIRE DU PERSONNEL OCCASIONNEL"]:
            r = center_p.add_run(line + "\n")
            r.bold = True
            r.underline = True
            r.font.size = Pt(11)

        right_p = right_cell.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_before = Pt(0)
        right_p.paragraph_format.space_after = Pt(0)
        rr = right_p.add_run("ANNEXE : 9……….\n")
        rr.bold = True
        rr.font.size = Pt(9)

        add_centered_title(doc_ref, "ROLE DES JOURNEES D'OUVRIERS EMPLOYES", font_size_pt=11)

        info_tbl = doc_ref.add_table(rows=1, cols=5)
        info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_docx_table_borders(info_tbl)
        info_widths = [usable_w_mm - (8 + 26 + 8 + 26), 8, 26, 8, 26]
        set_docx_table_fixed_layout(info_tbl, total_width_mm=usable_w_mm, col_widths_mm=info_widths)

        info_cells = info_tbl.rows[0].cells
        for cell in info_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        format_cell_text(
            info_cells[0],
            f"NOM DU REGISSEUR : {data['regisseur']}",
            align=WD_ALIGN_PARAGRAPH.LEFT,
            bold=True,
            size_pt=9,
        )
        format_cell_text(info_cells[1], "DU :", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=9)
        format_cell_text(info_cells[3], "AU :", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=9)

        for date_cell, date_text in ((info_cells[2], data["startDate"]), (info_cells[4], data["endDate"])):
            format_cell_text(date_cell, date_text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size_pt=9)
            for edge in ("top", "left", "bottom", "right"):
                set_docx_cell_border(date_cell, **{edge: {"val": "single", "sz": 10, "color": "000000"}})

        travaux_p = doc_ref.add_paragraph()
        travaux_p.paragraph_format.space_before = Pt(1)
        travaux_p.paragraph_format.space_after = Pt(1)
        travaux_p.add_run("(2)TRAVAUX DIVERS A LA COMMUNE OULED NACEUR").bold = True

        refs_p = doc_ref.add_paragraph()
        refs_p.paragraph_format.space_before = Pt(0)
        refs_p.paragraph_format.space_after = Pt(0)
        ref_parts = [
            ("ANNEE :", year),
            ("CHAP :", str(reference.get("chapitre") or reference.get("chap") or "")),
            ("Art :", str(reference.get("article") or reference.get("art") or "")),
            ("Prog :", str(reference.get("programme") or reference.get("prog") or "")),
            ("Proj :", str(reference.get("projet") or reference.get("proj") or "")),
            ("Ligne :", str(reference.get("ligne") or "")),
        ]
        for label, value in ref_parts:
            rl = refs_p.add_run(f"{label} ")
            rl.bold = True
            rv = refs_p.add_run(f"{value}    ")
            rv.bold = False
            rl.font.size = Pt(9)
            rv.font.size = Pt(9)

        somme_p = doc_ref.add_paragraph()
        somme_p.paragraph_format.space_before = Pt(0)
        somme_p.paragraph_format.space_after = Pt(2)
        rs = somme_p.add_run("SOMME A PAYER : ")
        rs.bold = True
        rs.font.size = Pt(9)
        rv = somme_p.add_run(fmt_amount(data["totalNet"], decimal_comma=decimal_comma))
        rv.font.size = Pt(9)

    def add_role_table(
        doc_ref,
        rows: List[Dict[str, Any]],
        totals: Dict[str, Any],
        is_continuation: bool = False,
        report_totals: Optional[Dict[str, Any]] = None,
        force_blank_rows: bool = True,
    ) -> None:
        table = doc_ref.add_table(rows=1, cols=9)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_docx_table_fixed_layout(table, total_width_mm=usable_w_mm, col_widths_mm=col_widths_mm)
        # Equivalent to CSS-like generous body padding (approx 8px vertical / 10px horizontal).
        set_docx_table_cell_margins(table, top_mm=2.8, bottom_mm=2.8, left_mm=3.5, right_mm=3.5)

        header = table.rows[0]
        header.height = Mm(layout.table_header_h_mm)
        header.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        header_labels = [
            "N° DESP.\nD'ATTACH.",
            "PRENOMS ET NOMS",
            "EMPLOIS",
            "NOMBRE DE JOURNEES",
            "PRIX DE LA JOURNEE",
            "BRUT A PAYER",
            "PRELEVEMENT I.G.R %6",
            "NET A PAYER",
            "N° DE LA C.I.N\nET SIGNATURE",
        ]
        rotated_header_cols = {0, 2, 3, 4, 5, 6}
        for idx, label in enumerate(header_labels):
            format_cell_text(
                header.cells[idx],
                label,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True,
                size_pt=8 if idx in (0, 3, 4, 6, 8) else 9,
            )
            header.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx in rotated_header_cols:
                set_docx_cell_text_direction(header.cells[idx], "btLr")

        slot_count = continuation_worker_rows if is_continuation else page1_worker_rows
        slot_count = max(slot_count, len(rows))
        base_start = page1_worker_rows + 1 if is_continuation else 1
        if rows:
            base_start = int(rows[0].get("rowNo") or base_start)

        rendered_rows = [dict(r) for r in rows]
        if force_blank_rows:
            while len(rendered_rows) < slot_count:
                rendered_rows.append({"rowNo": base_start + len(rendered_rows), "_blank": True})

        def add_summary_row(label: str, totals_row: Dict[str, Any]) -> None:
            row = table.add_row()
            row.height = Mm(layout.table_row_h_mm)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for c in row.cells:
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            label_cell = row.cells[0].merge(row.cells[2])
            format_cell_text(label_cell, label, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=9)
            format_cell_text(row.cells[3], str(int(round(to_float(totals_row.get("totalDays"), 0)))), bold=True, size_pt=9)
            format_cell_text(row.cells[4], "", bold=True, size_pt=9)
            format_cell_text(row.cells[5], fmt_amount(totals_row.get("totalGross"), decimal_comma=decimal_comma), bold=True, size_pt=9)
            format_cell_text(row.cells[6], fmt_amount(totals_row.get("totalDeduction"), decimal_comma=decimal_comma), bold=True, size_pt=9)
            format_cell_text(row.cells[7], fmt_amount(totals_row.get("totalNet"), decimal_comma=decimal_comma), bold=True, size_pt=9)
            format_cell_text(row.cells[8], "", bold=True, size_pt=9)

        if is_continuation:
            add_summary_row("REPORT :", report_totals or {"totalDays": 0, "totalGross": 0, "totalDeduction": 0, "totalNet": 0})

        for row_data in rendered_rows:
            row = table.add_row()
            row.height = Mm(layout.table_row_h_mm)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for c in row.cells:
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            is_blank = bool(row_data.get("_blank"))
            type_code = str(row_data.get("type") or "").strip()
            type_display = "O.S" if type_code == "OS" else ("O.N.S" if type_code else "")

            format_cell_text(row.cells[0], str(row_data.get("rowNo") or ""), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=9)
            format_cell_text(
                row.cells[1],
                "" if is_blank else str(row_data.get("nom_prenom") or ""),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                size_pt=9,
            )
            format_cell_text(row.cells[2], "" if is_blank else type_display, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=9)
            format_cell_text(
                row.cells[3],
                "" if is_blank else str(int(round(to_float(row_data.get("daysWorked"), 0)))),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size_pt=9,
            )
            format_cell_text(
                row.cells[4],
                "" if is_blank else fmt_amount(row_data.get("salaire_journalier"), decimal_comma=decimal_comma),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size_pt=9,
            )
            format_cell_text(
                row.cells[5],
                "" if is_blank else fmt_amount(row_data.get("grossSalary"), decimal_comma=decimal_comma),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size_pt=9,
            )
            format_cell_text(
                row.cells[6],
                "" if is_blank else fmt_amount(row_data.get("deduction"), decimal_comma=decimal_comma),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size_pt=9,
            )
            format_cell_text(
                row.cells[7],
                "" if is_blank else fmt_amount(row_data.get("netSalary"), decimal_comma=decimal_comma),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size_pt=9,
            )

            cin_cell = row.cells[8]
            cin_cell.text = ""
            if not is_blank:
                cin_value = str(row_data.get("cin") or "").strip()
                cin_validite = format_doc_date(row_data.get("cin_validite"))
                p1 = cin_cell.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p1.paragraph_format.space_before = Pt(0)
                p1.paragraph_format.space_after = Pt(0)
                r1 = p1.add_run(f"CIN N°: {cin_value}" if cin_value else "")
                r1.font.size = Pt(8)
                p2 = cin_cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0)
                r2 = p2.add_run(f"AU : {cin_validite}" if cin_validite else "AU :")
                r2.font.size = Pt(8)

        add_summary_row("TOTAL :", totals)

        for row in table.rows:
            for cell in row.cells:
                set_docx_cell_border(
                    cell,
                    top={"val": "single", "sz": 10, "color": "000000"},
                    left={"val": "single", "sz": 10, "color": "000000"},
                    bottom={"val": "single", "sz": 10, "color": "000000"},
                    right={"val": "single", "sz": 10, "color": "000000"},
                )

    def add_declaration_and_signatures(doc_ref, totals_in_words: str, dates: Dict[str, str]) -> None:
        # Keep a visible top margin between the table and the declaration/signature area.
        top_gap_tbl = doc_ref.add_table(rows=1, cols=1)
        remove_docx_table_borders(top_gap_tbl)
        set_docx_table_fixed_layout(top_gap_tbl, total_width_mm=usable_w_mm, col_widths_mm=[usable_w_mm])
        top_gap_row = top_gap_tbl.rows[0]
        top_gap_row.height = Mm(layout.post_block_margin_top_mm)
        top_gap_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        top_tbl = doc_ref.add_table(rows=1, cols=2)
        top_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_docx_table_borders(top_tbl)
        set_docx_table_fixed_layout(top_tbl, total_width_mm=usable_w_mm, col_widths_mm=[left_block_mm, right_block_mm])

        left_cell = top_tbl.rows[0].cells[0]
        right_cell = top_tbl.rows[0].cells[1]
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        format_cell_text(left_cell, "NOUS SOUSSIGNONS :", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size_pt=9)
        p_mr1 = left_cell.add_paragraph()
        p_mr1.paragraph_format.space_before = Pt(0)
        p_mr1.paragraph_format.space_after = Pt(0)
        p_mr1.paragraph_format.left_indent = Mm(layout.post_indent_mm)
        r_mr1 = p_mr1.add_run("Mr")
        r_mr1.bold = True
        r_mr1.font.size = Pt(9)

        p_blank = left_cell.add_paragraph("\u00a0")
        p_blank.paragraph_format.left_indent = Mm(layout.post_indent_mm)
        p_blank.paragraph_format.space_before = Pt(0)
        p_blank.paragraph_format.space_after = Pt(0)
        p_blank.runs[0].font.size = Pt(9)

        p_mr2 = left_cell.add_paragraph()
        p_mr2.paragraph_format.space_before = Pt(8)
        p_mr2.paragraph_format.space_after = Pt(0)
        p_mr2.paragraph_format.left_indent = Mm(layout.post_indent_mm)
        r_mr2 = p_mr2.add_run("Mr")
        r_mr2.bold = True
        r_mr2.font.size = Pt(9)

        p_cert = left_cell.add_paragraph()
        p_cert.paragraph_format.space_before = Pt(3)
        p_cert.paragraph_format.space_after = Pt(0)
        r_cert = p_cert.add_run("CERTIFIONS QUE LES SIEURS :")
        r_cert.bold = True
        r_cert.font.size = Pt(9)

        p_phrase = left_cell.add_paragraph()
        p_phrase.paragraph_format.space_before = Pt(1)
        p_phrase.paragraph_format.space_after = Pt(0)
        p_phrase.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_phrase = p_phrase.add_run("Portés au présent Role ont été payés en notre opposition de leurs signatures")
        r_phrase.font.size = Pt(9)

        format_cell_text(
            right_cell,
            "LE PRESENT ROLE S'ELEVANT A LA SOMME DE :",
            align=WD_ALIGN_PARAGRAPH.LEFT,
            bold=False,
            size_pt=9,
        )
        p_words = right_cell.add_paragraph()
        p_words.paragraph_format.space_before = Pt(1)
        p_words.paragraph_format.space_after = Pt(0)
        p_words.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rw = p_words.add_run(totals_in_words)
        rw.bold = True
        rw.font.size = Pt(9)

        center_line_1 = doc_ref.add_paragraph()
        center_line_1.paragraph_format.space_before = Pt(2)
        center_line_1.paragraph_format.space_after = Pt(0)
        center_line_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc1 = center_line_1.add_run("DRESSE ET CERTIFIE CONFORME AUX ATTACHEMENTS")
        rc1.bold = True
        rc1.font.size = Pt(9)

        center_line_2 = doc_ref.add_paragraph()
        center_line_2.paragraph_format.space_before = Pt(0)
        center_line_2.paragraph_format.space_after = Pt(0)
        center_line_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc2 = center_line_2.add_run(f"A OULED NACEUR LE : {dates.get('document_date') or ''}")
        rc2.font.size = Pt(9)

        spacer_tbl = doc_ref.add_table(rows=1, cols=1)
        remove_docx_table_borders(spacer_tbl)
        set_docx_table_fixed_layout(spacer_tbl, total_width_mm=usable_w_mm, col_widths_mm=[usable_w_mm])
        spacer_row = spacer_tbl.rows[0]
        spacer_row.height = Mm(layout.post_bottom_padding_mm)
        spacer_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        bottom_tbl = doc_ref.add_table(rows=1, cols=2)
        bottom_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_docx_table_borders(bottom_tbl)
        set_docx_table_fixed_layout(bottom_tbl, total_width_mm=usable_w_mm, col_widths_mm=[left_block_mm, right_block_mm])

        left_bottom = bottom_tbl.rows[0].cells[0]
        right_bottom = bottom_tbl.rows[0].cells[1]
        format_cell_text(
            left_bottom,
            f"PAYER PAR Moi Le : {dates.get('pay_date') or ''}",
            align=WD_ALIGN_PARAGRAPH.LEFT,
            bold=True,
            size_pt=9,
        )
        p_reg = left_bottom.add_paragraph()
        p_reg.paragraph_format.space_before = Pt(4)
        p_reg.paragraph_format.space_after = Pt(0)
        p_reg.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = p_reg.add_run("LE REGISSEUR DE DEPENSES")
        rr.bold = True
        rr.underline = True
        rr.font.size = Pt(10)

        format_cell_text(
            right_bottom,
            "LE PRESIDENT DU CONSEIL",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            size_pt=10,
            underline=True,
        )

        reserve_tbl = doc_ref.add_table(rows=1, cols=1)
        remove_docx_table_borders(reserve_tbl)
        set_docx_table_fixed_layout(reserve_tbl, total_width_mm=usable_w_mm, col_widths_mm=[usable_w_mm])
        reserve_row = reserve_tbl.rows[0]
        reserve_row.height = Mm(layout.post_signature_reserve_mm)
        reserve_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    def add_table_top_spacing(doc_ref, spacing_mm: float = 10.0) -> None:
        spacer_tbl = doc_ref.add_table(rows=1, cols=1)
        remove_docx_table_borders(spacer_tbl)
        set_docx_table_fixed_layout(spacer_tbl, total_width_mm=usable_w_mm, col_widths_mm=[usable_w_mm])
        spacer_row = spacer_tbl.rows[0]
        spacer_row.height = Mm(spacing_mm)
        spacer_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    for idx, sec in enumerate(sections):
        sec_workers = sec.get("workers") or []
        prepared_rows: List[Dict[str, Any]] = []
        for i, worker in enumerate(sec_workers):
            row = dict(worker)
            row["rowNo"] = i + 1
            prepared_rows.append(row)

        page1_rows = prepared_rows[:page1_worker_rows]
        page2_rows = prepared_rows[page1_worker_rows:]

        page1_partial = sum_workers(page1_rows)
        section_totals = normalize_totals(sec, sum_workers(prepared_rows))

        start_date = format_doc_date(sec.get("startDate"))
        end_date = format_doc_date(sec.get("endDate"))
        document_date = format_doc_date(sec.get("documentDate")) or end_date
        pay_date = format_doc_date(sec.get("payDate")) or document_date

        add_role_header_page1(
            doc,
            {
                "regisseur": regisseur,
                "startDate": start_date,
                "endDate": end_date,
                "totalNet": section_totals["totalNet"],
            },
        )
        add_table_top_spacing(doc, spacing_mm=10.0)
        add_role_table(
            doc,
            page1_rows,
            page1_partial,
            is_continuation=False,
            report_totals=None,
            force_blank_rows=True,
        )
        doc.add_page_break()

        add_role_table(
            doc,
            page2_rows,
            section_totals,
            is_continuation=True,
            report_totals=page1_partial,
            force_blank_rows=True,
        )

        total_net = round2(section_totals.get("totalNet"))
        words = number_to_words_fr(total_net).upper()
        cents = int(round((total_net - math.floor(total_net)) * 100.0))
        totals_words = f"{words} DHS {str(cents).rjust(2, '0')} CTS."

        add_declaration_and_signatures(
            doc,
            totals_words,
            {"document_date": document_date, "pay_date": pay_date},
        )

        if idx < len(sections) - 1:
            doc.add_page_break()

    doc.save(docx_path)


def main() -> None:
    payload = parse_input_json()
    output_dir = str(payload.get("outputDir") or "").strip()
    if not output_dir:
        raise ValueError("Missing required field: outputDir")

    ensure_dir(output_dir)

    safe_start = safe_filename_part(payload.get("safeStart") or payload.get("periodStart") or "")
    safe_end = safe_filename_part(payload.get("safeEnd") or payload.get("periodEnd") or "")

    docx_name = f"role_journees_{safe_start}_{safe_end}.docx"
    docx_path = os.path.join(output_dir, docx_name)

    try:
        draw_role_docx(payload, docx_path)
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Missing Python dependency for Word generation. Install: python-docx"
        ) from exc

    result = {"docxFileName": docx_name, "docxFilePath": docx_path}

    sys.stdout.write(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        error = {"success": False, "message": str(exc)}
        sys.stdout.write(json.dumps(error, ensure_ascii=False))
        sys.exit(1)
