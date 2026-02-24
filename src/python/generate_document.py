#!/usr/bin/env python3
import calendar
import json
import math
import os
import re
import sys
from datetime import date, datetime
from typing import Any, Dict, List, Optional, Tuple


RCAR_RATE = 0.06


def safe_filename_part(value: Any) -> str:
    text = str(value or "").strip()
    text = text.replace("\\", "-").replace("/", "-")
    text = re.sub(r"[^0-9A-Za-z._-]+", "_", text)
    text = text.strip("._-")
    return text or "unknown"


def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def parse_input_json() -> Dict[str, Any]:
    raw = sys.stdin.buffer.read()
    if not raw:
        raise ValueError("No JSON input received on stdin")
    try:
        return json.loads(raw.decode("utf-8"))
    except Exception as exc:
        raise ValueError(f"Invalid JSON input: {exc}") from exc


def parse_date(value: Any) -> Optional[date]:
    if not value:
        return None
    if isinstance(value, date) and not isinstance(value, datetime):
        return value
    if isinstance(value, datetime):
        return value.date()

    text = str(value).strip()
    if not text:
        return None

    if "/" in text:
        parts = [p.strip() for p in text.split("/")]
        if len(parts) == 3:
            try:
                d = int(parts[0])
                m = int(parts[1])
                y = int(parts[2])
                return date(y, m, d)
            except Exception:
                return None

    try:
        if "T" in text:
            return datetime.fromisoformat(text.replace("Z", "+00:00")).date()
        return datetime.fromisoformat(text).date()
    except Exception:
        return None


def calculate_age_at(birth: Any, ref: date) -> Optional[int]:
    b = parse_date(birth)
    if not b:
        return None
    age = ref.year - b.year
    if (ref.month, ref.day) < (b.month, b.day):
        age -= 1
    return age


def round2(value: Any) -> float:
    return float(f"{float(value or 0):.2f}")


def calculate_igr_deduction(date_naissance: Any, ref_date: date, gross_salary: float, age_limit: int) -> float:
    age = calculate_age_at(date_naissance, ref_date)
    if age is None or age <= int(age_limit):
        return round2(float(gross_salary) * RCAR_RATE)
    return 0.0


def calculate_net_salary(date_naissance: Any, ref_date: date, gross_salary: float, age_limit: int) -> float:
    return round2(float(gross_salary) - calculate_igr_deduction(date_naissance, ref_date, gross_salary, age_limit))


def month_start_end(year: int, month: int) -> Tuple[date, date]:
    end_day = calendar.monthrange(year, month)[1]
    return date(year, month, 1), date(year, month, end_day)


def fmt_amount(amount: float) -> str:
    return f"{float(amount or 0):.2f}"


def fmt_amount_receipt(amount: Any) -> str:
    try:
        value = float(amount or 0)
    except Exception:
        value = 0.0
    return f"{value:,.2f}".replace(",", " ")


def number_to_words_fr(num: float) -> str:
    whole_part = int(math.floor(float(num or 0)))

    units = [
        "Zero",
        "Un",
        "Deux",
        "Trois",
        "Quatre",
        "Cinq",
        "Six",
        "Sept",
        "Huit",
        "Neuf",
        "Dix",
        "Onze",
        "Douze",
        "Treize",
        "Quatorze",
        "Quinze",
        "Seize",
        "Dix Sept",
        "Dix Huit",
        "Dix Neuf",
    ]
    tens = ["", "", "Vingt", "Trente", "Quarante", "Cinquante", "Soixante"]

    def two_digits(n: int) -> str:
        if n < 20:
            return units[n]
        if n < 70:
            ten = n // 10
            unit = n % 10
            return tens[ten] if unit == 0 else f"{tens[ten]} {units[unit]}"
        if n < 80:
            return f"Soixante {two_digits(n - 60)}"
        if n == 80:
            return "Quatre Vingt"
        return f"Quatre Vingt {two_digits(n - 80)}"

    def three_digits(n: int) -> str:
        hundred = n // 100
        rest = n % 100
        if hundred == 0:
            hundred_part = ""
        elif hundred == 1:
            hundred_part = "Cent"
        else:
            hundred_part = f"{units[hundred]} Cent"

        if rest == 0:
            return hundred_part
        if not hundred_part:
            return two_digits(rest)
        return f"{hundred_part} {two_digits(rest)}"

    def to_words(n: int) -> str:
        if n == 0:
            return "Zero"
        parts: List[str] = []
        millions = n // 1_000_000
        thousands = (n % 1_000_000) // 1000
        rest = n % 1000

        if millions > 0:
            parts.append(f"{three_digits(millions)} Million" + ("s" if millions > 1 else ""))
        if thousands > 0:
            parts.append("Mille" if thousands == 1 else f"{three_digits(thousands)} Mille")
        if rest > 0:
            parts.append(three_digits(rest))
        return " ".join(parts)

    return to_words(whole_part)


DOC_FILE_BASE_MAP = {
    "demande-autorisation": "Demande_Autorisation_Paiement",
    "certificat-paiement": "Certificat_Paiement",
    "certificat-paiement-combined": "Certificat_Paiement",
    "ordre-paiement": "Ordre_Paiement",
    "mandat-paiement": "Mandat_Paiement",
    "depense-regie-recapitulatif": "Depense_Regie_Recapitulatif",
    "reference-values": "Valeurs_Reference",
    "bordereau": "Bordereau",
    "recu-combined": "Recu",
    "rcar-salariale": "RCAR",
    "rcar-patronale": "RCAR_PATRONALE",
}

DOC_TITLE_MAP = {
    "demande-autorisation": "DEMANDE D'AUTORISATION DE PAIEMENT",
    "certificat-paiement": "CERTIFICAT DE PAIEMENT",
    "certificat-paiement-combined": "CERTIFICAT DE PAIEMENT",
    "ordre-paiement": "ORDRE DE PAIEMENT",
    "mandat-paiement": "MANDAT DE PAIEMENT",
    "depense-regie-recapitulatif": "DEPENSE EN REGIE (RECAPITULATIF)",
    "reference-values": "VALEURS DE REFERENCE",
    "bordereau": "BORDEREAU",
    "recu-combined": "RECU",
    "rcar-salariale": "ETAT DE VERSEMENT RCAR (COTISATION SALARIALE)",
    "rcar-patronale": "ETAT DE VERSEMENT RCAR (COTISATION PATRONALE)",
}


def build_workers_financial_rows(
    report_rows: List[Dict[str, Any]],
    ref_date: date,
    age_limit: int,
) -> Tuple[List[Dict[str, Any]], Dict[str, float]]:
    rows_out: List[Dict[str, Any]] = []
    totals = {"days": 0.0, "gross": 0.0, "deduction": 0.0, "net": 0.0}

    for w in report_rows:
        days = int(w.get("days_worked") or w.get("total_days") or 0)
        gross = float(w.get("amount") or w.get("total_amount") or 0.0)
        if days <= 0 or gross <= 0:
            continue

        deduction = calculate_igr_deduction(w.get("date_naissance"), ref_date, gross, age_limit)
        net = round2(gross - deduction)

        totals["days"] += days
        totals["gross"] += gross
        totals["deduction"] += deduction
        totals["net"] += net

        rows_out.append(
            {
                "nom_prenom": str(w.get("nom_prenom") or ""),
                "cin": str(w.get("cin") or ""),
                "type": str(w.get("type") or ""),
                "days": days,
                "gross": round2(gross),
                "deduction": round2(deduction),
                "net": round2(net),
            }
        )

    totals = {k: round2(v) for k, v in totals.items()}
    return rows_out, totals


def build_docx_filename(document_type: str, payload: Dict[str, Any]) -> str:
    base = DOC_FILE_BASE_MAP.get(document_type, safe_filename_part(document_type))
    year = payload.get("year")
    month = payload.get("month")
    quarter = payload.get("quarter")
    options = payload.get("options") or {}

    if document_type == "bordereau":
        number = safe_filename_part(options.get("bordereauNumber") or month or "12")
        if year and month:
            return f"{base}_{number}_{year}.docx"
        return f"{base}_{number}.docx"

    if document_type in ("rcar-salariale", "rcar-patronale") and year and quarter:
        months = ((payload.get("report") or {}).get("period") or {}).get("months") or []
        if months:
            period_from = date(int(year), int(months[0]), 1)
            period_to = date(int(year), int(months[-1]), calendar.monthrange(int(year), int(months[-1]))[1])
        else:
            period_from = date(int(year), 1, 1)
            period_to = date(int(year), 1, 1)
        from_safe = f"{period_from.day:02d}-{period_from.month:02d}-{period_from.year}"
        to_safe = f"{period_to.day:02d}-{period_to.month:02d}-{period_to.year}"
        return f"{base}_{year}_T{quarter}_{from_safe}_{to_safe}.docx"

    if year and month:
        return f"{base}_{str(month).zfill(2)}_{year}.docx"
    if year and quarter:
        return f"{base}_T{quarter}_{year}.docx"
    return f"{base}.docx"


def set_page_margins(section, *, top_mm: float, bottom_mm: float, left_mm: float, right_mm: float) -> None:
    from docx.shared import Mm

    section.top_margin = Mm(top_mm)
    section.bottom_margin = Mm(bottom_mm)
    section.left_margin = Mm(left_mm)
    section.right_margin = Mm(right_mm)


def set_default_font(doc, *, font_name: str, font_size_pt: float) -> None:
    from docx.shared import Pt

    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size_pt)


def add_centered_title(doc, text: str, *, font_size_pt: float) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size_pt)


def add_table_with_widths(doc, *, cols: int, col_widths_mm: List[float]):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Mm

    t = doc.add_table(rows=1, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    for i, w in enumerate(col_widths_mm):
        if i < len(t.columns):
            t.columns[i].width = Mm(w)
    return t


def first_paragraph(doc):
    # python-docx creates an empty paragraph by default; reuse it to avoid an extra blank line at the top.
    try:
        if getattr(doc, "paragraphs", None) and doc.paragraphs:
            return doc.paragraphs[0]
    except Exception:
        pass
    return doc.add_paragraph()


def remove_leading_empty_paragraph(doc) -> None:
    try:
        if not getattr(doc, "paragraphs", None) or not doc.paragraphs:
            return
        p = doc.paragraphs[0]
        if (p.text or "").strip():
            return
        element = p._p
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    except Exception:
        return

def remove_docx_table_borders(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "nil")


def generate_recu_docx(payload: Dict[str, Any], docx_path: str) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt
    except ModuleNotFoundError as exc:
        raise RuntimeError("Missing Python dependency for Word generation. Install: python-docx") from exc

    document_type = str(payload.get("documentType") or "").strip()
    if document_type != "recu-combined":
        raise ValueError(f"Unsupported document type for recu generator: {document_type}")

    year = int(payload.get("year") or 0)
    month = int(payload.get("month") or 0)
    if year <= 0 or month <= 0 or month > 12:
        raise ValueError("Missing required fields for recu: year, month")
    report = payload.get("report") or {}
    report_rows = report.get("rows") or []
    options = payload.get("options") or {}

    age_limit = int(options.get("rcarAgeLimit") or 60)
    _, end_date = month_start_end(year, month)

    _, totals = build_workers_financial_rows(report_rows, end_date, age_limit)
    total_net = round2(totals.get("net") or 0)

    amount_digits = fmt_amount_receipt(total_net)
    words = number_to_words_fr(total_net)
    whole_part = int(math.floor(float(total_net or 0)))
    cents = int(round((float(total_net or 0) - float(whole_part)) * 100.0))
    cents_str = str(cents).rjust(2, "0")
    amount_words = f"{words} dhs {cents_str} Cts"

    regisseur_name = str(options.get("regisseurName") or "MAJDA TAKNOUTI")
    regisseur_cin = str(options.get("regisseurCin") or "I 528862")
    regisseur_cin_valid = str(options.get("regisseurCinValidUntil") or "30/09/2030")
    province = str(options.get("provinceName") or "FQUIH BEN SALAH")
    commune = str(options.get("communeName") or "OULED NACEUR")
    city = str(options.get("cityName") or "Ouled Naceur")

    doc = Document()
    section0 = doc.sections[0]
    section0.page_width = Mm(210)
    section0.page_height = Mm(297)
    set_page_margins(section0, top_mm=15, bottom_mm=15, left_mm=20, right_mm=20)
    set_default_font(doc, font_name="Times New Roman", font_size_pt=11)

    # Header (top-left)
    p_header = first_paragraph(doc)
    p_header.text = ""
    p_header.paragraph_format.space_before = Pt(0)
    p_header.paragraph_format.space_after = Pt(60)
    r1 = p_header.add_run("ROYAUME DU MAROC\nMINISTERE DE L'INTERIEUR\n")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p_header.add_run(f"PROVINCE DE {province}\n")
    r2.bold = True
    r2.font.size = Pt(10)
    r3 = p_header.add_run(f"COMMUNE {commune}")
    r3.bold = True
    r3.underline = True
    r3.font.size = Pt(10)

    # Title (center)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(18)
    rt1 = p_title.add_run("Recu")
    rt1.bold = True
    rt1.underline = True
    rt1.font.size = Pt(12)
    rt2 = p_title.add_run("   N°   ")
    rt2.bold = True
    rt2.font.size = Pt(12)
    rt3 = p_title.add_run(f"{month}/{year}")
    rt3.bold = True
    rt3.font.size = Pt(12)

    # Body
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(12)
    p1.add_run("Je soussignée Mme : ")
    rn = p1.add_run(regisseur_name)
    rn.bold = True
    p1.add_run("    CIN  :  ")
    rc = p1.add_run(regisseur_cin)
    rc.bold = True
    p1.add_run(f"  valable au  {regisseur_cin_valid} Régisseur de dépenses titulaire")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(12)
    p2.add_run("Reconnaît avoir reçu la somme de :  ")
    ra = p2.add_run(amount_digits)
    ra.bold = True
    p2.add_run(f" ({amount_words}).")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(6)
    p3.paragraph_format.space_after = Pt(18)
    p3.add_run(f"{city} Le : " + "." * 45)

    # Signatures (top part)
    usable_w_mm = 210 - 20 - 20
    sig = add_table_with_widths(doc, cols=2, col_widths_mm=[usable_w_mm * 0.5, usable_w_mm * 0.5])
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_docx_table_borders(sig)
    for cell in sig.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    sig.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig.rows[0].cells[0].paragraphs[0].add_run("Percepteur de souk sebt")
    sig.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.rows[0].cells[1].paragraphs[0].add_run("Le Régisseur")

    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(75)

    # DEMANDE DE FONDS
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(18)
    rd = p4.add_run("DEMANDE DE FONDS")
    rd.bold = True
    rd.underline = True
    rd.font.size = Pt(14)

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_before = Pt(0)
    p5.paragraph_format.space_after = Pt(10)
    p5.add_run("Pour la date du : " + "." * 45 + " Le Régisseur  demande d'être approvisionné de la somme de :")

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_before = Pt(0)
    p6.paragraph_format.space_after = Pt(10)
    p6.add_run(
        f"{amount_digits} ({amount_words}). Qui sera retiré à la caisse du comptable de\u00A0rattachement."
    )

    p7 = doc.add_paragraph()
    p7.paragraph_format.space_before = Pt(0)
    p7.paragraph_format.space_after = Pt(24)
    p7.add_run(
        "Qui sera versé au compte courant postal n° "
        + "." * 38
        + " dont l'intitulé est : salaire du personnel occasionnel."
    )

    gap2 = doc.add_paragraph()
    gap2.paragraph_format.space_before = Pt(0)
    gap2.paragraph_format.space_after = Pt(30)

    p8 = doc.add_paragraph()
    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p8.paragraph_format.space_before = Pt(0)
    p8.paragraph_format.space_after = Pt(18)
    p8.add_run(f"A {city} le : " + "." * 25)

    sig2 = add_table_with_widths(doc, cols=2, col_widths_mm=[usable_w_mm * 0.5, usable_w_mm * 0.5])
    sig2.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_docx_table_borders(sig2)
    for cell in sig2.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    sig2.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig2.rows[0].cells[0].paragraphs[0].add_run("Visa du Président")
    sig2.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig2.rows[0].cells[1].paragraphs[0].add_run("Le Régisseur de dépenses")

    doc.save(docx_path)


def generate_demande_autorisation_docx(payload: Dict[str, Any], docx_path: str) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt
    except ModuleNotFoundError as exc:
        raise RuntimeError("Missing Python dependency for Word generation. Install: python-docx") from exc

    document_type = str(payload.get("documentType") or "").strip()
    if document_type != "demande-autorisation":
        raise ValueError(f"Unsupported document type for demande-autorisation generator: {document_type}")

    year = int(payload.get("year") or 0)
    month = int(payload.get("month") or 0)
    if year <= 0 or month <= 0 or month > 12:
        raise ValueError("Missing required fields for demande-autorisation: year, month")

    report = payload.get("report") or {}
    report_rows = report.get("rows") or []
    options = payload.get("options") or {}

    age_limit = int(options.get("rcarAgeLimit") or 60)
    _, end_date = month_start_end(year, month)
    _, totals = build_workers_financial_rows(report_rows, end_date, age_limit)
    total_net = round2(totals.get("net") or 0)

    amount_digits = fmt_amount_receipt(total_net)
    words = number_to_words_fr(total_net)
    whole_part = int(math.floor(float(total_net or 0)))
    cents = int(round((float(total_net or 0) - float(whole_part)) * 100.0))
    cents_str = str(cents).rjust(2, "0")
    amount_words = f"{words} dhs {cents_str} Cts"

    chap = str(options.get("chap") or "10").strip()
    art = str(options.get("art") or "20").strip()
    prog = str(options.get("prog") or "20").strip()
    proj = str(options.get("proj") or "10").strip()
    ligne = str(options.get("ligne") or "14").strip()

    regisseur_name = str(options.get("regisseurName") or "MAJDA TAKNOUTI")
    province = str(options.get("provinceName") or "FQUIH BEN SALAH")
    commune = str(options.get("communeName") or "OULED NACEUR")
    city = str(options.get("cityName") or "Ouled Naceur")

    doc = Document()
    section0 = doc.sections[0]
    section0.page_width = Mm(210)
    section0.page_height = Mm(297)
    set_page_margins(section0, top_mm=20, bottom_mm=20, left_mm=20, right_mm=20)
    set_default_font(doc, font_name="Times New Roman", font_size_pt=11)

    # Header (top-left)
    p_header = first_paragraph(doc)
    p_header.text = ""
    p_header.paragraph_format.space_before = Pt(0)
    p_header.paragraph_format.space_after = Pt(18)
    rh1 = p_header.add_run("ROYAUME DU MAROC\nMINISTERE DE L'INTERIEUR\n")
    rh1.bold = True
    rh1.font.size = Pt(10)
    rh2 = p_header.add_run(f"PROVINCE DE {province}\n")
    rh2.bold = True
    rh2.font.size = Pt(10)
    rh3 = p_header.add_run(f"COMMUNE {commune}")
    rh3.bold = True
    rh3.underline = True
    rh3.font.size = Pt(10)

    # Title (center)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(14)
    rt = p_title.add_run(f"DEMANDE D'AUTORISATION DE PAIEMENT N° {month}/{year}")
    rt.bold = True
    rt.underline = True
    rt.font.size = Pt(11)

    # Body
    p1 = doc.add_paragraph("Vu L’arrêté conjointe de création de régie de dépenses auprès de la commune Ouled Naceur.")
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(10)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(10)
    p2.add_run("Mme : ")
    rn = p2.add_run(regisseur_name)
    rn.bold = True
    p2.add_run(
        " demande l’autorisation de payer par voie de régie de dépenses imputées sur les rubriques budgétaires citée ci après désignée :"
    )

    def add_ref_line(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Mm(6)
        p.add_run(f"{label:<10} :  {value}")

    add_ref_line("Chapitre", chap)
    add_ref_line("Article", art)
    add_ref_line("Programme", prog)
    add_ref_line("Projet", proj)
    add_ref_line("Ligne", ligne)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Table (budget line)
    usable_w_mm = 210 - 40  # 20mm left + 20mm right
    widths = [
        usable_w_mm * 0.08,
        usable_w_mm * 0.10,
        usable_w_mm * 0.10,
        usable_w_mm * 0.09,
        usable_w_mm * 0.10,
        usable_w_mm * 0.33,
        usable_w_mm * 0.20,
    ]
    table = add_table_with_widths(doc, cols=7, col_widths_mm=widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["CHAP", "ART", "PROG", "PROJ", "LIGNE", "INTITULE", "MONTANT"]
    hdr = table.rows[0]
    for cell in hdr.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for i, text in enumerate(headers):
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)

    row = table.add_row()
    row.height = Mm(35)
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    values = [chap, art, prog, proj, ligne, "Salaire du personnel\noccasionnel", amount_digits]
    for i, value in enumerate(values):
        p = row.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(value))
        r.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    p3 = doc.add_paragraph(f"Arrêté la présente demande à la somme de :  {amount_digits} ({amount_words}).")
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(12)

    p4 = doc.add_paragraph(f"A {city} Le : " + "." * 35)
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(24)

    # Signatures
    sig = add_table_with_widths(doc, cols=2, col_widths_mm=[usable_w_mm * 0.5, usable_w_mm * 0.5])
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_docx_table_borders(sig)
    for cell in sig.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    sig.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.rows[0].cells[0].paragraphs[0].add_run("Le Président")
    sig.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.rows[0].cells[1].paragraphs[0].add_run("Le régisseur")

    doc.save(docx_path)


def generate_certificat_paiement_docx(payload: Dict[str, Any], docx_path: str) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt
    except ModuleNotFoundError as exc:
        raise RuntimeError("Missing Python dependency for Word generation. Install: python-docx") from exc

    document_type = str(payload.get("documentType") or "").strip()
    if document_type not in ("certificat-paiement", "certificat-paiement-combined"):
        raise ValueError(f"Unsupported document type for certificat-paiement generator: {document_type}")

    year = int(payload.get("year") or 0)
    month = int(payload.get("month") or 0)
    if year <= 0 or month <= 0 or month > 12:
        raise ValueError("Missing required fields for certificat-paiement: year, month")

    report = payload.get("report") or {}
    report_rows = report.get("rows") or []
    options = payload.get("options") or {}

    age_limit = int(options.get("rcarAgeLimit") or 60)
    _, end_date = month_start_end(year, month)
    _, totals = build_workers_financial_rows(report_rows, end_date, age_limit)
    total_net = round2(totals.get("net") or 0)

    amount_digits = fmt_amount_receipt(total_net)
    words = number_to_words_fr(total_net)
    whole_part = int(math.floor(float(total_net or 0)))
    cents = int(round((float(total_net or 0) - float(whole_part)) * 100.0))
    cents_str = str(cents).rjust(2, "0")
    amount_words = f"{words} dhs {cents_str} Cts"

    regisseur_name = str(options.get("regisseurName") or "MAJDA TAKNOUTI").strip() or "MAJDA TAKNOUTI"
    decision_number = str(options.get("decisionNumber") or "").strip()
    decision_date = str(options.get("decisionDate") or "").strip()
    if not decision_number or not decision_date:
        raise ValueError(
            "Missing CERTIFICAT DE PAIEMENT decision reference in configuration "
            "(decision number/date)."
        )
    document_date = str(options.get("documentDate") or "........................").strip() or "........................"
    exercise_year = str(options.get("exerciseYear") or year).strip() or str(year)

    chap = str(options.get("chap") or "10").strip()
    art = str(options.get("art") or "20").strip()
    prog = str(options.get("prog") or "20").strip()
    proj = str(options.get("proj") or "10").strip()
    ligne = str(options.get("ligne") or "14").strip()

    province = str(options.get("provinceName") or "FQUIH BEN SALAH")
    commune = str(options.get("communeName") or "OULED NACEUR")
    city = str(options.get("cityName") or "Ouled Naceur")

    doc = Document()
    section0 = doc.sections[0]
    section0.page_width = Mm(210)
    section0.page_height = Mm(297)
    set_page_margins(section0, top_mm=20, bottom_mm=20, left_mm=20, right_mm=20)
    set_default_font(doc, font_name="Times New Roman", font_size_pt=11)

    # Header (top-left)
    p_header = first_paragraph(doc)
    p_header.text = ""
    p_header.paragraph_format.space_before = Pt(0)
    p_header.paragraph_format.space_after = Pt(28)
    r1 = p_header.add_run("ROYAUME DU MAROC\nMINISTERE DE L'INTERIEUR\n")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p_header.add_run(f"PROVINCE DE {province}\n")
    r2.bold = True
    r2.font.size = Pt(10)
    r3 = p_header.add_run(f"COMMUNE {commune}")
    r3.bold = True
    r3.underline = True
    r3.font.size = Pt(10)

    # Title (centered, underlined)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(26)
    rt = p_title.add_run("CERTIFICAT DE PAIEMENT")
    rt.bold = True
    rt.underline = True
    rt.font.size = Pt(14)

    # Body (justify, continuous paragraph)
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body.paragraph_format.space_before = Pt(0)
    p_body.paragraph_format.space_after = Pt(30)
    p_body.paragraph_format.line_spacing = 1.15

    p_body.add_run("Le Président de la commune Ouled Naceur : Vu la décision N° ")
    p_body.add_run(decision_number)
    p_body.add_run(f" en date du {decision_date}. Relative à la nomination de Mme : ")
    rn = p_body.add_run(regisseur_name)
    rn.bold = True
    p_body.add_run(
        " régisseur de dépenses; considérant une avance de : "
        f"{amount_digits} ({amount_words}). "
        "Non justifiées est nécessaire au régisseur de dépenses pour paiement des ouvriers : "
        "Travaux divers à la commune Ouled Naceur. "
        "Certifie qu’il est à Mme : Le régisseur à Ouled Naceur de payer sur le budget de l’exercice "
        f"{exercise_year} "
        f"Chap : {chap}  Art : {art}  Prog : {prog}  Proj : {proj}  Ligne : {ligne} . "
        "relative au salaire du personnel occasionnel à la somme de "
        f"{amount_digits} ({amount_words})."
    )

    # Date + signature
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after = Pt(28)
    p_date.add_run(f"A {city} Le : {document_date}")

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig.paragraph_format.space_before = Pt(0)
    p_sig.paragraph_format.space_after = Pt(0)
    p_sig.add_run("Le Président")

    doc.save(docx_path)


def generate_ordre_paiement_docx(payload: Dict[str, Any], docx_path: str) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt
    except ModuleNotFoundError as exc:
        raise RuntimeError("Missing Python dependency for Word generation. Install: python-docx") from exc

    document_type = str(payload.get("documentType") or "").strip()
    if document_type != "ordre-paiement":
        raise ValueError(f"Unsupported document type for ordre-paiement generator: {document_type}")

    year = int(payload.get("year") or 0)
    month = int(payload.get("month") or 0)
    if year <= 0 or month <= 0 or month > 12:
        raise ValueError("Missing required fields for ordre-paiement: year, month")

    report = payload.get("report") or {}
    report_rows = report.get("rows") or []
    options = payload.get("options") or {}

    age_limit = int(options.get("rcarAgeLimit") or 60)
    _, end_date = month_start_end(year, month)
    _, totals = build_workers_financial_rows(report_rows, end_date, age_limit)
    total_net = round2(totals.get("net") or 0)

    amount_digits = fmt_amount_receipt(total_net)
    words = number_to_words_fr(total_net)
    whole_part = int(math.floor(float(total_net or 0)))
    cents = int(round((float(total_net or 0) - float(whole_part)) * 100.0))
    cents_str = str(cents).rjust(2, "0")
    amount_words = f"{words} dhs {cents_str} Cts"

    chap = str(options.get("chap") or "10").strip()
    art = str(options.get("art") or "20").strip()
    prog = str(options.get("prog") or "20").strip()
    proj = str(options.get("proj") or "10").strip()
    ligne = str(options.get("ligne") or "14").strip()

    regisseur_name = str(options.get("regisseurName") or "MAJDA TAKNOUTI").strip() or "MAJDA TAKNOUTI"
    document_date = str(options.get("documentDate") or "........................").strip() or "........................"
    province = str(options.get("provinceName") or "FQUIH BEN SALAH")
    commune = str(options.get("communeName") or "OULED NACEUR")
    city = str(options.get("cityName") or "Ouled Naceur")

    period_from = date(year, month, 1)
    period_to = end_date
    period_from_s = f"{period_from.day:02d}/{period_from.month:02d}/{period_from.year}"
    period_to_s = f"{period_to.day:02d}/{period_to.month:02d}/{period_to.year}"

    doc = Document()
    section0 = doc.sections[0]
    section0.page_width = Mm(210)
    section0.page_height = Mm(297)
    set_page_margins(section0, top_mm=20, bottom_mm=20, left_mm=20, right_mm=20)
    set_default_font(doc, font_name="Times New Roman", font_size_pt=11)

    # Header (top-left)
    p_header = first_paragraph(doc)
    p_header.text = ""
    p_header.paragraph_format.space_before = Pt(0)
    p_header.paragraph_format.space_after = Pt(55)
    r1 = p_header.add_run("ROYAUME DU MAROC\nMINISTERE DE L'INTERIEUR\n")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p_header.add_run(f"PROVINCE DE {province}\n")
    r2.bold = True
    r2.font.size = Pt(10)
    r3 = p_header.add_run(f"COMMUNE {commune}")
    r3.bold = True
    r3.underline = True
    r3.font.size = Pt(10)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(26)
    rt = p_title.add_run("ORDRE DE PAIEMENT")
    rt.bold = True
    rt.underline = True
    rt.font.size = Pt(14)

    # Body (justify)
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body.paragraph_format.space_before = Pt(0)
    p_body.paragraph_format.space_after = Pt(26)
    p_body.paragraph_format.line_spacing = 1.15

    p_body.add_run(
        f"Sur les crédits disponibles de la rubrique budgétaire : Chap : {chap} Art : {art} Prog : {prog} "
        f"Proj : {proj} Ligne : {ligne} . "
        "relative au salaire du personnel occasionnel "
        "L’ordre est donc donné à Mme "
    )
    rn = p_body.add_run(regisseur_name)
    rn.bold = True
    p_body.add_run(
        f" régisseur de dépenses à la commune {city} de payer la somme de "
        f"{amount_digits} ({amount_words}). "
        f"Comme salaire du personnel occasionnel à la commune {city} ."
    )

    # Period line
    p_period = doc.add_paragraph(f"Période du :  {period_from_s} au  {period_to_s}.")
    p_period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_period.paragraph_format.space_before = Pt(0)
    p_period.paragraph_format.space_after = Pt(26)

    # Date line
    p_date = doc.add_paragraph(f"{city} Le : {document_date}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after = Pt(32)

    # Signature
    p_sig = doc.add_paragraph("Le Président")
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig.paragraph_format.space_before = Pt(0)
    p_sig.paragraph_format.space_after = Pt(0)

    doc.save(docx_path)


def generate_mandat_paiement_docx(payload: Dict[str, Any], docx_path: str) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt
    except ModuleNotFoundError as exc:
        raise RuntimeError("Missing Python dependency for Word generation. Install: python-docx") from exc

    document_type = str(payload.get("documentType") or "").strip()
    if document_type != "mandat-paiement":
        raise ValueError(f"Unsupported document type for mandat-paiement generator: {document_type}")

    year = int(payload.get("year") or 0)
    month = int(payload.get("month") or 0)
    if year <= 0 or month <= 0 or month > 12:
        raise ValueError("Missing required fields for mandat-paiement: year, month")

    report = payload.get("report") or {}
    report_rows = report.get("rows") or []
    options = payload.get("options") or {}

    age_limit = int(options.get("rcarAgeLimit") or 60)
    _, end_date = month_start_end(year, month)
    _, totals = build_workers_financial_rows(report_rows, end_date, age_limit)
    total_net = round2(totals.get("net") or 0)

    amount_digits = fmt_amount_receipt(total_net)
    words = number_to_words_fr(total_net)
    whole_part = int(math.floor(float(total_net or 0)))
    cents = int(round((float(total_net or 0) - float(whole_part)) * 100.0))
    cents_str = str(cents).rjust(2, "0")
    amount_words = f"{words} dhs {cents_str} Cts"

    exercise_year = str(options.get("exerciseYear") or year).strip() or str(year)
    mandat_number = str(options.get("mandatNumber") or f"{month}/{year}").strip()
    bordereau_number = str(options.get("bordereauNumber") or "").strip()
    document_date = str(options.get("documentDate") or "........................").strip() or "........................"

    chap = str(options.get("chap") or "10").strip()
    art = str(options.get("art") or "20").strip()
    prog = str(options.get("prog") or "20").strip()
    proj = str(options.get("proj") or "10").strip()
    ligne = str(options.get("ligne") or "14").strip()

    regisseur_name = str(options.get("regisseurName") or "MAJDA TAKNOUTI").strip() or "MAJDA TAKNOUTI"
    province = str(options.get("provinceName") or "FQUIH BEN SALAH")
    commune = str(options.get("communeName") or "OULED NACEUR")
    city = str(options.get("cityName") or "Ouled Naceur")

    period_from = date(year, month, 1)
    period_to = end_date
    period_from_s = f"{period_from.day:02d}/{period_from.month:02d}/{period_from.year}"
    period_to_s = f"{period_to.day:02d}/{period_to.month:02d}/{period_to.year}"

    doc = Document()
    remove_leading_empty_paragraph(doc)
    section0 = doc.sections[0]
    section0.page_width = Mm(210)
    section0.page_height = Mm(297)
    set_page_margins(section0, top_mm=20, bottom_mm=20, left_mm=20, right_mm=15)
    set_default_font(doc, font_name="Times New Roman", font_size_pt=11)

    usable_w_mm = 210 - 20 - 15

    # Header block (two columns: left header, right exercise/mandat)
    header_tbl = add_table_with_widths(doc, cols=2, col_widths_mm=[usable_w_mm * 0.60, usable_w_mm * 0.40])
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_tbl.autofit = False
    remove_docx_table_borders(header_tbl)

    for cell in header_tbl.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    c_left = header_tbl.rows[0].cells[0]
    c_left.text = ""
    p_left = c_left.paragraphs[0]
    p_left.paragraph_format.space_before = Pt(0)
    p_left.paragraph_format.space_after = Pt(0)
    p_left.paragraph_format.line_spacing = 1.0

    r = p_left.add_run("ROYAUME DU MAROC")
    r.bold = True
    r.font.size = Pt(10)
    r.add_break()
    r2 = p_left.add_run("MINISTERE DE L'INTERIEUR")
    r2.bold = True
    r2.font.size = Pt(10)
    r2.add_break()
    r3 = p_left.add_run(f"PROVINCE DE {province}")
    r3.bold = True
    r3.font.size = Pt(10)
    r3.add_break()
    r4 = p_left.add_run(f"COMMUNE {commune}")
    r4.bold = True
    r4.underline = True
    r4.font.size = Pt(10)

    c_right = header_tbl.rows[0].cells[1]
    c_right.text = ""
    p_right = c_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_before = Pt(0)
    p_right.paragraph_format.space_after = Pt(0)
    right_text = f"EXERCICE : {exercise_year}   MANDAT N° : {mandat_number}".rstrip()
    rr = p_right.add_run(right_text)
    rr.font.size = Pt(10)

    # Spacing after header
    p_gap = doc.add_paragraph()
    p_gap.paragraph_format.space_before = Pt(0)
    p_gap.paragraph_format.space_after = Pt(18)

    # Budget refs
    p_budget = doc.add_paragraph(f"Chap: {chap}, Art : {art}, Prog : {prog}, Proj : {proj}, ligne :{ligne}")
    p_budget.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_budget.paragraph_format.space_before = Pt(0)
    p_budget.paragraph_format.space_after = Pt(8)
    for run in p_budget.runs:
        run.font.size = Pt(9)

    # Salary line
    p_salary = doc.add_paragraph("SALAIRE DU PERSONNEL OCCASIONNEL")
    p_salary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_salary.paragraph_format.space_before = Pt(0)
    p_salary.paragraph_format.space_after = Pt(10)
    rs = p_salary.runs[0]
    rs.bold = True
    rs.font.size = Pt(10)

    # Title line
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    rt = p_title.add_run("MANDAT DE PAIEMENT DELIVRE PAR NOUS ORDONNATEUR")
    rt.bold = True
    rt.underline = True
    rt.font.size = Pt(11)

    # Main table
    col_widths_mm = [
        usable_w_mm * 0.36,
        usable_w_mm * 0.26,
        usable_w_mm * 0.12,
        usable_w_mm * 0.26,
    ]
    table = add_table_with_widths(doc, cols=4, col_widths_mm=col_widths_mm)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = [
        "Prénom et Nom\nQualité et Résidence de la partie prenante",
        "Objet de la Depense",
        "Montant",
        "Pièces produites\nà l'appui du Mandat",
    ]
    hdr = table.rows[0]
    hdr.height = Mm(8)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.text = ""
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(7)

    row = table.add_row()
    row.height = Mm(50)
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Column 1
    c0 = row.cells[0]
    c0.text = ""
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.line_spacing = 1.0
    r_name = p0.add_run(regisseur_name)
    r_name.bold = True
    r_name.font.size = Pt(10)
    r_name.add_break()
    r_role = p0.add_run("REGISSEUR")
    r_role.underline = True
    r_role.font.size = Pt(9)
    r_role.add_break()
    r_city = p0.add_run(f"A LA COMMUNE {commune}")
    r_city.font.size = Pt(8)

    # Column 2
    c1 = row.cells[1]
    c1.text = ""
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.line_spacing = 1.0
    r_obj1 = p1.add_run("SALAIRE")
    r_obj1.bold = True
    r_obj1.font.size = Pt(9)
    r_obj1.add_break()
    r_obj2 = p1.add_run("DU PERSONNEL OCCASIONNEL")
    r_obj2.bold = True
    r_obj2.font.size = Pt(9)

    # Column 3
    c2 = row.cells[2]
    c2.text = ""
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.0
    r_amt = p2.add_run(amount_digits)
    r_amt.font.size = Pt(10)

    # Column 4
    c3 = row.cells[3]
    c3.text = ""
    p3 = c3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.line_spacing = 1.0
    bordereau_line = f"BORDEREAU N° {month}/{year}"
    pieces_text = "\n".join(
        [
            "CERTIFICAT DE PAIEMENT",
            "FEUILLE D'ATTACHEMENT",
            "ORDRE DE PAIEMENT",
            "ROLE DE JOURNEE",
            bordereau_line,
        ]
    )
    r_pieces = p3.add_run(pieces_text)
    r_pieces.font.size = Pt(8)

    # After table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

    p_sum = doc.add_paragraph(f"LA SOMME DE : {amount_digits} ({amount_words}).")
    p_sum.paragraph_format.space_before = Pt(0)
    p_sum.paragraph_format.space_after = Pt(6)
    if p_sum.runs:
        p_sum.runs[0].font.size = Pt(9)

    # Payable text on one line, with the date on the next line.
    p_lp = doc.add_paragraph()
    p_lp.paragraph_format.space_before = Pt(0)
    p_lp.paragraph_format.space_after = Pt(0)
    p_lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_lp = p_lp.add_run(
        "A LAQUELLE S'ELEVE LE PRESENT MANDAT EST PAYABLE PAR LE PERCEPTEUR DE SOUK SEBT TRESORIER DE LA COMMUNE"
    )
    r_lp.font.size = Pt(8)

    p_rp = doc.add_paragraph()
    p_rp.paragraph_format.space_before = Pt(0)
    p_rp.paragraph_format.space_after = Pt(0)
    p_rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_rp = p_rp.add_run(f"Le : {document_date}")
    r_rp.font.size = Pt(8)

    # Spacer before signatures
    p_gap2 = doc.add_paragraph()
    p_gap2.paragraph_format.space_before = Pt(0)
    p_gap2.paragraph_format.space_after = Pt(36)

    # Signature blocks
    sig_tbl = add_table_with_widths(doc, cols=2, col_widths_mm=[usable_w_mm * 0.50, usable_w_mm * 0.50])
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_tbl.autofit = False
    remove_docx_table_borders(sig_tbl)

    # First row
    sig_tbl.rows[0].cells[0].text = ""
    p_sl = sig_tbl.rows[0].cells[0].paragraphs[0]
    p_sl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sl.paragraph_format.space_before = Pt(0)
    p_sl.paragraph_format.space_after = Pt(0)
    r_sl = p_sl.add_run("VU BON A PAYER PAR LE\nPERCEPTEUR DE SOUK SEBT\nTRESORIER DE LA COMMUNE")
    r_sl.font.size = Pt(8)

    sig_tbl.rows[0].cells[1].text = ""
    p_sr = sig_tbl.rows[0].cells[1].paragraphs[0]
    p_sr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sr.paragraph_format.space_before = Pt(0)
    p_sr.paragraph_format.space_after = Pt(0)
    r_sr = p_sr.add_run("L'ORDONNATEUR DE LA COMMUNE")
    r_sr.font.size = Pt(8)

    # Second row (right acquit)
    row2 = sig_tbl.add_row()
    row2.cells[0].text = ""
    row2.cells[1].text = ""
    p_ar = row2.cells[1].paragraphs[0]
    p_ar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ar.paragraph_format.space_before = Pt(0)
    p_ar.paragraph_format.space_after = Pt(0)
    r_ar = p_ar.add_run("POUR ACQUIT DE LA SOMME CI-DESSUS\nLA PARTIE PRENANTE")
    r_ar.font.size = Pt(8)

    doc.save(docx_path)


def fmt_amount_fr_comma(amount: Any) -> str:
    try:
        value = float(amount or 0)
    except Exception:
        value = 0.0
    return f"{value:,.2f}".replace(",", " ").replace(".", ",")


def amount_to_words_dhs_cents(amount: Any) -> str:
    try:
        value = float(amount or 0)
    except Exception:
        value = 0.0
    whole_part = int(math.floor(value))
    cents = int(round((value - float(whole_part)) * 100.0))
    cents_str = str(cents).rjust(2, "0")
    words = number_to_words_fr(value)
    return f"{words} dhs {cents_str} Cts"


def calculate_range_net_total(
    report_rows: List[Dict[str, Any]],
    *,
    year: int,
    month: int,
    start_day: int,
    end_day: int,
    age_limit: int,
) -> float:
    ref_date = date(int(year), int(month), int(end_day))
    total_cents = 0

    for w in report_rows:
        presence_days = w.get("presenceDays") or w.get("presence_days") or []
        if not isinstance(presence_days, list):
            presence_days = []
        daily_salary = float(w.get("salaire_journalier") or w.get("dailySalary") or w.get("daily_salary") or 0.0)
        if daily_salary <= 0:
            continue
        days_in_range = 0
        for d in presence_days:
            try:
                day_num = int(d)
            except Exception:
                continue
            if start_day <= day_num <= end_day:
                days_in_range += 1
        if days_in_range <= 0:
            continue

        gross = round2(days_in_range * daily_salary)
        deduction = calculate_igr_deduction(w.get("date_naissance"), ref_date, gross, age_limit)
        net = round2(gross - deduction)
        total_cents += int(round(net * 100))

    return total_cents / 100.0


def generate_bordereau_docx(payload: Dict[str, Any], docx_path: str) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt
    except ModuleNotFoundError as exc:
        raise RuntimeError("Missing Python dependency for Word generation. Install: python-docx") from exc

    document_type = str(payload.get("documentType") or "").strip()
    if document_type != "bordereau":
        raise ValueError(f"Unsupported document type for bordereau generator: {document_type}")

    year = int(payload.get("year") or 0)
    month = int(payload.get("month") or 0)
    if year <= 0 or month <= 0 or month > 12:
        raise ValueError("Missing required fields for bordereau: year, month")

    report = payload.get("report") or {}
    report_rows = report.get("rows") or []
    options = payload.get("options") or {}

    age_limit = int(options.get("rcarAgeLimit") or 60)
    last_day = calendar.monthrange(year, month)[1]

    # Ensure we have some payroll data for the period (combined range)
    combined_net = calculate_range_net_total(
        report_rows,
        year=year,
        month=month,
        start_day=1,
        end_day=last_day,
        age_limit=age_limit,
    )
    if combined_net <= 0:
        raise ValueError("Aucune donnée de paie pour cette période")

    bordereau_number = str(options.get("bordereauNumber") or "12").strip() or "12"
    document_date = str(options.get("documentDate") or "........................").strip() or "........................"
    chap = str(options.get("chap") or "10").strip()
    art = str(options.get("art") or "20").strip()
    prog = str(options.get("prog") or "20").strip()
    proj = str(options.get("proj") or "10").strip()
    ligne = str(options.get("ligne") or "14").strip()

    regisseur_name = str(options.get("regisseurName") or "Majda Taknouti").strip() or "Majda Taknouti"
    report_previous = float(options.get("reportPreviousBordereau") or 0)
    rejected_amount = float(options.get("rejectedAmount") or 0)

    province = str(options.get("provinceName") or "FQUIH BEN SALAH")
    commune = str(options.get("communeName") or "OULED NACEUR")
    city = str(options.get("cityName") or "Ouled Naceur")

    doc = Document()
    remove_leading_empty_paragraph(doc)

    section0 = doc.sections[0]
    section0.page_width = Mm(210)
    section0.page_height = Mm(297)
    set_page_margins(section0, top_mm=15, bottom_mm=15, left_mm=15, right_mm=10)
    set_default_font(doc, font_name="Times New Roman", font_size_pt=11)

    usable_w_mm = 210 - 15 - 10

    def add_header_block() -> None:
        hdr = add_table_with_widths(doc, cols=2, col_widths_mm=[usable_w_mm * 0.70, usable_w_mm * 0.30])
        hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr.autofit = False
        remove_docx_table_borders(hdr)

        left_cell = hdr.rows[0].cells[0]
        right_cell = hdr.rows[0].cells[1]
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        left_cell.text = ""
        p = left_cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        for line, underline in [
            ("ROYAUME DU MAROC", False),
            ("MINISTERE DE L'INTERIEUR", False),
            (f"PROVINCE DE {province}", False),
            (f"COMMUNE {commune}", True),
        ]:
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(9)
            if underline:
                r.underline = True
            r.add_break()

        right_cell.text = ""
        pr = right_cell.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pr.paragraph_format.space_before = Pt(0)
        pr.paragraph_format.space_after = Pt(0)
        rr1 = pr.add_run("D.216")
        rr1.bold = True
        rr1.font.size = Pt(9)
        rr1.add_break()
        rr2 = pr.add_run("ANNEXE 19")
        rr2.bold = True
        rr2.font.size = Pt(9)
        rr2.add_break()
        rr3 = pr.add_run("Titre............................")
        rr3.font.size = Pt(9)

    def add_centered_line(text: str, *, bold: bool = False, underline: bool = False, size: int = 10, after: int = 0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = bool(bold)
        r.underline = bool(underline)
        r.font.size = Pt(size)

    def add_para(text: str, *, align: str = "left", size: int = 9, after: int = 0, line_spacing: float = 1.15):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = line_spacing
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "justify":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.size = Pt(size)
        return p

    def add_dots(cell, count: int, *, align=WD_ALIGN_PARAGRAPH.LEFT, size_pt: int = 8):
        for _ in range(count):
            p = cell.add_paragraph("." * 45)
            p.alignment = align
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.size = Pt(size_pt)

    def mm_to_twips(value_mm: float) -> int:
        return int(round(float(value_mm) * 72.0 / 25.4 * 20.0))

    def set_table_fixed_layout(table, col_widths_mm: List[float]) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        table.autofit = False
        tbl_pr = table._tbl.tblPr

        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")

        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(mm_to_twips(usable_w_mm)))

        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_ind.set(qn("w:w"), "0")

        for col_idx, width_mm in enumerate(col_widths_mm):
            if col_idx >= len(table.columns):
                break
            table.columns[col_idx].width = Mm(width_mm)
            for row in table.rows:
                if col_idx < len(row.cells):
                    row.cells[col_idx].width = Mm(width_mm)

    def set_table_borders(table, sz: int = 8) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tbl_pr = table._tbl.tblPr
        tbl_borders = tbl_pr.find(qn("w:tblBorders"))
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = tbl_borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tbl_borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), str(sz))
            element.set(qn("w:color"), "000000")

    def set_cell_borders(cell, *, top: int = 0, left: int = 0, bottom: int = 0, right: int = 0) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)

        for edge, sz in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
            if sz <= 0:
                continue
            edge_tag = qn(f"w:{edge}")
            edge_el = tc_borders.find(edge_tag)
            if edge_el is None:
                edge_el = OxmlElement(f"w:{edge}")
                tc_borders.append(edge_el)
            edge_el.set(qn("w:val"), "single")
            edge_el.set(qn("w:sz"), str(sz))
            edge_el.set(qn("w:color"), "000000")

    def set_table_cell_margins(table, *, top_mm: float, bottom_mm: float, left_mm: float, right_mm: float) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tbl_pr = table._tbl.tblPr
        cell_mar = tbl_pr.find(qn("w:tblCellMar"))
        if cell_mar is None:
            cell_mar = OxmlElement("w:tblCellMar")
            tbl_pr.append(cell_mar)

        for edge, mm_val in (("top", top_mm), ("bottom", bottom_mm), ("left", left_mm), ("right", right_mm)):
            node = cell_mar.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                cell_mar.append(node)
            node.set(qn("w:w"), str(mm_to_twips(mm_val)))
            node.set(qn("w:type"), "dxa")

    ranges = [
        ("combined", 1, last_day),
    ]

    for idx, (_rid, start_day, end_day) in enumerate(ranges):
        range_net = calculate_range_net_total(
            report_rows,
            year=year,
            month=month,
            start_day=start_day,
            end_day=end_day,
            age_limit=age_limit,
        )
        bord_amount = fmt_amount_fr_comma(range_net)
        bord_words = amount_to_words_dhs_cents(range_net)

        admitted_raw = options.get("admittedAmount")
        if admitted_raw in (None, ""):
            admitted = range_net
        else:
            try:
                admitted = float(admitted_raw)
            except Exception:
                admitted = range_net
        admitted = max(round2(admitted), 0.0)
        admitted_amount = fmt_amount_fr_comma(admitted)
        rejected_amount_s = fmt_amount_fr_comma(rejected_amount)
        report_previous_s = fmt_amount_fr_comma(report_previous)
        total_general = report_previous + range_net
        total_general_s = fmt_amount_fr_comma(total_general)
        total_general_words = amount_to_words_dhs_cents(total_general)

        date_from = date(year, month, start_day)
        date_to = date(year, month, end_day)
        date_from_s = f"{date_from.day:02d}/{date_from.month:02d}/{date_from.year}"
        date_to_s = f"{date_to.day:02d}/{date_to.month:02d}/{date_to.year}"

        # --- Page 1 ---
        add_header_block()
        add_centered_line("DÉPENSE EN RÉGIE", bold=True, size=11, after=2)
        add_centered_line("Salaire du Personnel Occasionnel", underline=True, size=11, after=1)
        add_centered_line(
            f"Titre Chap: {chap} Art {art} Prog: {prog}, Pro: {proj}, ligne: {ligne}.",
            size=9,
            after=6,
        )

        # Regisseur block (left)
        p_reg = doc.add_paragraph()
        p_reg.paragraph_format.space_before = Pt(0)
        p_reg.paragraph_format.space_after = Pt(8)
        p_reg.paragraph_format.line_spacing = 1.25
        p_reg.add_run(f"Mme: {regisseur_name}\n").font.size = Pt(9)
        p_reg.add_run("Régisseur de dépenses\n").font.size = Pt(9)
        p_reg.add_run("Régie de dépenses auprès\n").font.size = Pt(9)
        p_reg.add_run(f"COMMUNE {commune}").font.size = Pt(9)

        add_centered_line(f"BORDEREAU N: {month}/{year}", bold=True, size=9, after=1)
        add_centered_line(f"du {date_from_s} AU {date_to_s}", bold=True, size=9, after=4)

        add_para(
            "des quittances et des pièces adressées à M(1): Le percepteur, par le soussigné,\n"
            "pour justifier l’emploi des fonds qui lui ont été remis par le percepteur",
            align="justify",
            size=10,
            after=2,
            line_spacing=1.15,
        )
        add_para(f"d’un montant de: {bord_amount} ({bord_words}).", align="justify", size=10, after=4, line_spacing=1.15)

        # Main table (page 1): keep fully on first page with tighter row sizing/padding.
        col_w = [
            usable_w_mm * 0.08,
            usable_w_mm * 0.25,
            usable_w_mm * 0.18,
            usable_w_mm * 0.15,
            usable_w_mm * 0.17,
            usable_w_mm * 0.17,
        ]
        t = add_table_with_widths(doc, cols=6, col_widths_mm=col_w)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(t, col_w)
        # Tighter padding prevents overflow while keeping readability.
        set_table_cell_margins(t, top_mm=0.4, bottom_mm=0.4, left_mm=1.0, right_mm=1.0)
        set_table_borders(t, sz=8)

        headers = [
            "Numéro des\nPièces",
            "Désignation des\nPièces",
            "Nature des\nDépenses",
            "Montant des\nparties prenantes",
            "Nom des\nparties prenantes",
            "Observations",
        ]
        hdr = t.rows[0]
        hdr.height = Mm(12)
        hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(8)

        row = t.add_row()
        row.height = Mm(104)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0

        row.cells[0].text = ""

        # Désignation
        c1 = row.cells[1]
        c1.text = ""
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run("Rôle de journée\nOrdre de paiement\nCertificat de paiement\nFeuille d’attachement")
        r1.font.size = Pt(8.5)

        # Nature
        c2 = row.cells[2]
        c2.text = ""
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run("Salaire du Personnel\nOccasionnel")
        r2.font.size = Pt(8.5)

        # Montant (right aligned)
        c3 = row.cells[3]
        c3.text = ""
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r3 = p3.add_run(bord_amount)
        r3.font.size = Pt(8.5)
        r3.bold = True

        # Nom des parties prenantes
        c4 = row.cells[4]
        c4.text = ""
        p4 = c4.paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r4 = p4.add_run("parties prenantes")
        r4.font.size = Pt(8.5)

        # Observations
        row.cells[5].text = ""

        # A REPORTER row inside the same main table
        rep_row = t.add_row()
        rep_row.height = Mm(7)
        rep_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in rep_row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        rep_row.cells[0].merge(rep_row.cells[2])
        rep_row.cells[0].text = ""
        p_rep = rep_row.cells[0].paragraphs[0]
        p_rep.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = p_rep.add_run("A REPORTER :")
        rr.bold = True
        rr.font.size = Pt(8.5)

        rep_row.cells[3].text = ""
        p_rep_amt = rep_row.cells[3].paragraphs[0]
        p_rep_amt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rra = p_rep_amt.add_run(bord_amount)
        rra.bold = True
        rra.font.size = Pt(8.5)

        # Force split exactly after A REPORTER
        doc.add_page_break()

        # --- Page 2 (continuation) ---
        t2 = add_table_with_widths(doc, cols=6, col_widths_mm=col_w)
        t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(t2, col_w)
        set_table_cell_margins(t2, top_mm=1.2, bottom_mm=1.2, left_mm=2.0, right_mm=2.0)
        set_table_borders(t2, sz=8)

        hdr2 = t2.rows[0]
        hdr2.height = Mm(16)
        hdr2.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for i, h in enumerate(headers):
            cell = hdr2.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(9)

        summary_rows = [
            ("Report :", report_previous_s),
            ("Total :", bord_amount),
            ("Total Général :", total_general_s),
        ]
        for label, value in summary_rows:
            sr = t2.add_row()
            sr.height = Mm(9)
            sr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in sr.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
            sr.cells[0].text = ""
            sr.cells[1].text = ""
            sr.cells[2].text = ""
            p_lbl = sr.cells[2].paragraphs[0]
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r_lbl = p_lbl.add_run(label)
            r_lbl.bold = True
            r_lbl.font.size = Pt(10)

            sr.cells[3].text = ""
            p_val = sr.cells[3].paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_val = p_val.add_run(value)
            r_val.bold = True
            r_val.font.size = Pt(10)
            sr.cells[4].text = ""
            sr.cells[5].text = ""

        gap_after_top = doc.add_paragraph()
        gap_after_top.paragraph_format.space_before = Pt(0)
        gap_after_top.paragraph_format.space_after = Pt(12)

        add_para(
            "Arrêté le présent bordereau, comprenant quittances et pièces à la somme total\n"
            f"de : {bord_words}.",
            align="justify",
            size=10,
            after=6,
            line_spacing=1.15,
        )
        add_para(f"A : {city}, Le : {document_date}", align="center", size=10, after=8)

        sig = add_table_with_widths(doc, cols=2, col_widths_mm=[usable_w_mm * 0.5, usable_w_mm * 0.5])
        sig.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(sig, [usable_w_mm * 0.5, usable_w_mm * 0.5])
        remove_docx_table_borders(sig)
        psl = sig.rows[0].cells[0].paragraphs[0]
        psr = sig.rows[0].cells[1].paragraphs[0]
        psl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        psr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        psl.paragraph_format.space_before = Pt(0)
        psl.paragraph_format.space_after = Pt(0)
        psr.paragraph_format.space_before = Pt(0)
        psr.paragraph_format.space_after = Pt(0)
        psl.add_run("Le Président").font.size = Pt(10)
        psr.add_run("Le Régisseur").font.size = Pt(10)

        sig_space = doc.add_paragraph()
        sig_space.paragraph_format.space_before = Pt(0)
        sig_space.paragraph_format.space_after = Pt(48)

        # Bottom section: unified table to keep left/right blocks perfectly aligned.
        bottom_w = [
            usable_w_mm * 0.56 * 0.12,
            usable_w_mm * 0.56 * 0.18,
            usable_w_mm * 0.56 * 0.70,
            usable_w_mm * 0.44,
        ]
        bottom_tbl = add_table_with_widths(doc, cols=4, col_widths_mm=bottom_w)
        bottom_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(bottom_tbl, bottom_w)
        set_table_borders(bottom_tbl, sz=8)
        set_table_cell_margins(bottom_tbl, top_mm=0.8, bottom_mm=0.8, left_mm=1.0, right_mm=1.0)

        b_title = bottom_tbl.rows[0]
        b_title.height = Mm(9)
        b_title.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for c in b_title.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        b_title.cells[0].merge(b_title.cells[2])
        b_title.cells[0].text = ""
        p_lt = b_title.cells[0].paragraphs[0]
        p_lt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lt = p_lt.add_run("Quittances et pièces rejetées")
        r_lt.bold = True
        r_lt.font.size = Pt(10)
        b_title.cells[3].text = ""
        p_rt = b_title.cells[3].paragraphs[0]
        p_rt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_rt = p_rt.add_run("Situation des dépenses")
        r_rt.bold = True
        r_rt.font.size = Pt(10)

        b_head = bottom_tbl.add_row()
        b_head.height = Mm(9)
        b_head.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for c in b_head.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        l_headers = ["N° d'ordre", "Montant", "Indications des pièces produites"]
        for i, text in enumerate(l_headers):
            b_head.cells[i].text = ""
            ph = b_head.cells[i].paragraphs[0]
            ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            rh = ph.add_run(text)
            rh.bold = True
            rh.font.size = Pt(9)
        b_head.cells[3].text = ""

        b_body = bottom_tbl.add_row()
        b_body.height = Mm(44)
        b_body.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for c in b_body.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            c.text = ""

        # Merge right-side content area to match left total height and keep a clean shared border.
        r_cell = b_head.cells[3].merge(b_body.cells[3])
        r_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        r_cell.text = ""
        situation_lines = [
            f"Montant du présent bordereau: {bord_amount}",
            f"Montant admis du présent bordereau: {admitted_amount}",
            f"Report du bordereau précédent: {report_previous_s}",
            f"Montant rejeté du présent bordereau: {rejected_amount_s}",
            f"Total Général: {total_general_s}",
        ]
        for i, line in enumerate(situation_lines):
            p = r_cell.paragraphs[0] if i == 0 else r_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            rr = p.add_run(line)
            rr.font.size = Pt(9)

        add_para(
            "Renvoi est fait régisseur désigné ci-dessus du présent bordereau été définitivement à la somme\n"
            f"de : {total_general_words}",
            align="justify",
            size=10,
            after=2,
        )
        add_para(
            "et des quittances et pièces non admises dont le montant est indiqué au tableau ci-dessus",
            align="justify",
            size=10,
            after=4,
        )
        add_para(f"Souk sebt, Le : …../…../{year}", align="right", size=10, after=1)
        add_para("le ....................(comptable-assignataire)", align="right", size=10, after=0)

        if idx < len(ranges) - 1:
            doc.add_page_break()

    doc.save(docx_path)


def _generate_rcar_docx(
    payload: Dict[str, Any],
    docx_path: str,
    *,
    expected_document_type: str,
    subtitle: str,
    prelevement_rate: float,
    justificatif_rg_key: str,
) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt
    except ModuleNotFoundError as exc:
        raise RuntimeError("Missing Python dependency for Word generation. Install: python-docx") from exc

    document_type = str(payload.get("documentType") or "").strip()
    if document_type != expected_document_type:
        raise ValueError(f"Unsupported document type for {expected_document_type} generator: {document_type}")

    report = payload.get("report") or {}
    report_rows = report.get("rows") or []
    options = payload.get("options") or {}

    def safe_int(value: Any) -> int:
        try:
            return int(str(value).strip())
        except Exception:
            return 0

    def safe_float(value: Any, default: float = 0.0) -> float:
        try:
            return float(value)
        except Exception:
            return float(default)

    def mm_to_twips(value_mm: float) -> int:
        return int(round(float(value_mm) * 72.0 / 25.4 * 20.0))

    def set_table_fixed_layout(table, col_widths_mm: List[float]) -> None:
        table.autofit = False
        tbl_pr = table._tbl.tblPr

        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")

        total_w_mm = sum(col_widths_mm)
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(mm_to_twips(total_w_mm)))

        for col_idx, width_mm in enumerate(col_widths_mm):
            if col_idx >= len(table.columns):
                break
            table.columns[col_idx].width = Mm(width_mm)
            for row in table.rows:
                if col_idx < len(row.cells):
                    row.cells[col_idx].width = Mm(width_mm)

    def set_table_borders(table, sz: int = 4) -> None:
        tbl_pr = table._tbl.tblPr
        tbl_borders = tbl_pr.find(qn("w:tblBorders"))
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge_el = tbl_borders.find(qn(f"w:{edge}"))
            if edge_el is None:
                edge_el = OxmlElement(f"w:{edge}")
                tbl_borders.append(edge_el)
            edge_el.set(qn("w:val"), "single")
            edge_el.set(qn("w:sz"), str(sz))
            edge_el.set(qn("w:color"), "000000")

    def set_table_cell_margins(table, *, top_mm: float, bottom_mm: float, left_mm: float, right_mm: float) -> None:
        tbl_pr = table._tbl.tblPr
        cell_mar = tbl_pr.find(qn("w:tblCellMar"))
        if cell_mar is None:
            cell_mar = OxmlElement("w:tblCellMar")
            tbl_pr.append(cell_mar)

        for edge, mm_val in (("top", top_mm), ("bottom", bottom_mm), ("left", left_mm), ("right", right_mm)):
            edge_el = cell_mar.find(qn(f"w:{edge}"))
            if edge_el is None:
                edge_el = OxmlElement(f"w:{edge}")
                cell_mar.append(edge_el)
            edge_el.set(qn("w:w"), str(mm_to_twips(mm_val)))
            edge_el.set(qn("w:type"), "dxa")

    def clear_cell(cell) -> None:
        cell.text = ""
        if cell.paragraphs:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    def quarter_months(q: int) -> List[int]:
        mapping = {1: [1, 2, 3], 2: [4, 5, 6], 3: [7, 8, 9], 4: [10, 11, 12]}
        return mapping.get(q, [])

    def quarter_months_label(months: List[int]) -> str:
        month_names = {
            1: "JANVIER",
            2: "FEVRIER",
            3: "MARS",
            4: "AVRIL",
            5: "MAI",
            6: "JUIN",
            7: "JUILLET",
            8: "AOUT",
            9: "SEPTEMBRE",
            10: "OCTOBRE",
            11: "NOVEMBRE",
            12: "DECEMBRE",
        }
        labels = [month_names.get(m, "") for m in months if m in month_names]
        return " ".join([x for x in labels if x]).strip()

    def fmt_amount_fr(value: Any, blank_zero: bool = False) -> str:
        try:
            v = float(value)
        except Exception:
            return ""
        if blank_zero and abs(v) < 1e-9:
            return ""
        return f"{v:,.2f}".replace(",", " ").replace(".", ",")

    def fmt_days(value: float) -> str:
        if abs(value - round(value)) < 1e-9:
            return str(int(round(value)))
        return f"{value:.2f}".replace(".", ",")

    def amount_words_upper(amount: float) -> str:
        value = abs(float(amount or 0))
        whole_part = int(math.floor(value))
        cents = int(round((value - whole_part) * 100))
        if cents == 100:
            whole_part += 1
            cents = 0
        words = number_to_words_fr(whole_part).upper()
        return f"{words} DHS {cents:02d} CTS"

    def render_digit_boxes(cell, digits: Any, *, boxes: int, box_width_mm: float = 4.5, font_size_pt: int = 9) -> None:
        clear_cell(cell)
        count = max(1, int(boxes))
        text = re.sub(r"\D", "", str(digits or ""))
        text = text[-count:]
        start = count - len(text)

        tbl = cell.add_table(rows=1, cols=count)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_fixed_layout(tbl, [box_width_mm] * count)
        set_table_borders(tbl, sz=4)
        set_table_cell_margins(tbl, top_mm=0.0, bottom_mm=0.0, left_mm=0.0, right_mm=0.0)

        row = tbl.rows[0]
        row.height = Mm(7.0)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for idx, c in enumerate(row.cells):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            ch = text[idx - start] if idx >= start else ""
            if ch:
                r = p.add_run(ch)
                r.font.size = Pt(font_size_pt)

    def render_amount_boxes(cell, amount: Optional[float], *, total_width_mm: float, boxes: int = 11) -> None:
        clear_cell(cell)
        if total_width_mm <= 8.0:
            total_width_mm = 8.0

        boxes_width = min(total_width_mm - 6.5, boxes * 4.5)
        if boxes_width < 4.5:
            boxes_width = max(4.5, total_width_mm - 6.5)
        decimals_width = max(4.5, total_width_mm - boxes_width)

        outer = cell.add_table(rows=1, cols=2)
        outer.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_fixed_layout(outer, [boxes_width, decimals_width])
        remove_docx_table_borders(outer)
        set_table_cell_margins(outer, top_mm=0.0, bottom_mm=0.0, left_mm=0.0, right_mm=0.0)

        box_cell = outer.rows[0].cells[0]
        dec_cell = outer.rows[0].cells[1]
        box_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        dec_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if amount is None:
            render_digit_boxes(box_cell, "", boxes=boxes, box_width_mm=4.5, font_size_pt=9)
            clear_cell(dec_cell)
            return

        value = abs(float(amount))
        integer = int(math.floor(value))
        cents = int(round((value - integer) * 100))
        if cents == 100:
            integer += 1
            cents = 0

        render_digit_boxes(box_cell, str(integer), boxes=boxes, box_width_mm=4.5, font_size_pt=9)

        clear_cell(dec_cell)
        p_dec = dec_cell.paragraphs[0]
        p_dec.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_dec.paragraph_format.space_before = Pt(0)
        p_dec.paragraph_format.space_after = Pt(0)
        r_dec = p_dec.add_run(f",{cents:02d}")
        r_dec.font.size = Pt(9)

    year = safe_int(payload.get("year"))
    quarter = safe_int(payload.get("quarter"))

    period_months_raw = ((report.get("period") or {}).get("months") or [])
    period_months = [safe_int(m) for m in period_months_raw if 1 <= safe_int(m) <= 12]
    if not period_months and quarter in (1, 2, 3, 4):
        period_months = quarter_months(quarter)
    if quarter not in (1, 2, 3, 4) and period_months:
        quarter = ((period_months[0] - 1) // 3) + 1

    period_from_s = ""
    period_to_s = ""
    if year > 0 and period_months:
        start_month = period_months[0]
        end_month = period_months[-1]
        try:
            d_from = date(year, start_month, 1)
            d_to = date(year, end_month, calendar.monthrange(year, end_month)[1])
            period_from_s = f"{d_from.day:02d}/{d_from.month:02d}/{d_from.year}"
            period_to_s = f"{d_to.day:02d}/{d_to.month:02d}/{d_to.year}"
        except Exception:
            period_from_s = ""
            period_to_s = ""

    table_rows = []
    total_days = 0.0
    total_brut = 0.0
    total_prelev = 0.0
    total_versement = 0.0

    for worker in report_rows:
        monthly_stats = worker.get("monthlyStats") if isinstance(worker.get("monthlyStats"), list) else []

        days_source = worker.get("total_days")
        if days_source in (None, ""):
            days_source = worker.get("days_worked")
        if days_source in (None, "") and monthly_stats:
            days_source = sum(safe_float(ms.get("days_worked"), 0.0) for ms in monthly_stats)
        days = safe_float(days_source, 0.0)

        brut_source = worker.get("total_amount")
        if brut_source in (None, ""):
            brut_source = worker.get("amount")
        if brut_source in (None, "") and monthly_stats:
            brut_source = sum(safe_float(ms.get("amount"), 0.0) for ms in monthly_stats)
        brut = round2(safe_float(brut_source, 0.0))

        if days <= 0 and brut <= 0:
            continue

        daily_price = safe_float(
            worker.get("salaire_journalier")
            or worker.get("dailySalary")
            or worker.get("daily_salary"),
            0.0,
        )
        if daily_price <= 0 and days > 0 and brut > 0:
            daily_price = round2(brut / days)

        prelev = round2(brut * prelevement_rate) if brut > 0 else 0.0
        versement = prelev

        table_rows.append(
            {
                "nom_prenom": str(worker.get("nom_prenom") or ""),
                "qualite": str(worker.get("type") or "").upper(),
                "days": days,
                "price": daily_price if daily_price > 0 else None,
                "brut": brut if brut > 0 else 0.0,
                "prelev": prelev,
                "versement": versement,
            }
        )

        total_days += days
        total_brut += brut
        total_prelev += prelev
        total_versement += versement

    total_brut = round2(total_brut)
    total_prelev = round2(total_prelev)
    total_versement = round2(total_versement)

    commune = str(options.get("communeName") or "OULED NACEUR").strip().upper()
    province = str(options.get("provinceName") or "FQUIH BEN SALAH").strip().upper()
    city = str(options.get("cityName") or "Ouled Naceur").strip() or "Ouled Naceur"
    arabic_header = str(options.get("rcarArabicLine") or "").strip()
    adhesion_number = re.sub(r"\D", "", str(options.get("rcarAdhesionNumber") or "35160001"))

    doc = Document()
    remove_leading_empty_paragraph(doc)
    section0 = doc.sections[0]
    section0.page_width = Mm(210)
    section0.page_height = Mm(297)
    set_page_margins(section0, top_mm=10, bottom_mm=10, left_mm=12, right_mm=12)
    set_default_font(doc, font_name="Times New Roman", font_size_pt=10)

    usable_w_mm = 210 - 12 - 12

    # ------------------------------
    # Page 1: Etat de versement
    # ------------------------------
    p_header = first_paragraph(doc)
    p_header.text = ""
    p_header.paragraph_format.space_before = Pt(0)
    p_header.paragraph_format.space_after = Pt(16)
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r0 = p_header.add_run("ROYAUME DU MAROC\n")
    r0.bold = True
    r0.font.size = Pt(10)
    r1 = p_header.add_run("MINISTERE DE L'INTERIEUR\n")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p_header.add_run(f"PROVINCE DE {province}\n")
    r2.bold = True
    r2.font.size = Pt(10)
    r3 = p_header.add_run(f"COMMUNE {commune}")
    r3.bold = True
    r3.font.size = Pt(10)

    p_title1 = doc.add_paragraph()
    p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title1.paragraph_format.space_before = Pt(0)
    p_title1.paragraph_format.space_after = Pt(0)
    rt1 = p_title1.add_run("ETAT DE VERSEMENT A LA (R.C.A.R)")
    rt1.bold = True
    rt1.underline = True
    rt1.font.size = Pt(14)

    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title2.paragraph_format.space_before = Pt(0)
    p_title2.paragraph_format.space_after = Pt(6)
    rt2 = p_title2.add_run(subtitle)
    rt2.bold = True
    rt2.underline = True
    rt2.font.size = Pt(14)

    p_period = doc.add_paragraph()
    p_period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_period.paragraph_format.space_before = Pt(0)
    p_period.paragraph_format.space_after = Pt(8)
    if period_from_s and period_to_s:
        period_text = f"Période du : {period_from_s} au {period_to_s}"
    else:
        period_text = "Période du : .......... au .........."
    rp = p_period.add_run(period_text)
    rp.font.size = Pt(12)

    col_w = [68.0, 16.0, 16.0, 22.0, 24.0, 20.0, 20.0]
    state_tbl = add_table_with_widths(doc, cols=7, col_widths_mm=col_w)
    state_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(state_tbl, col_w)
    set_table_borders(state_tbl, sz=4)
    set_table_cell_margins(state_tbl, top_mm=0.25, bottom_mm=0.25, left_mm=0.6, right_mm=0.6)

    headers = [
        "Nom et prénom",
        "Qualité",
        "Nbre de jours",
        "Prix de journée",
        "Brut à payer",
        f"Prélèvement {int(round(prelevement_rate * 100))}%",
        "Total de versement",
    ]

    hdr = state_tbl.rows[0]
    hdr.height = Mm(10.0)
    hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for idx, htxt in enumerate(headers):
        c = hdr.cells[idx]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(htxt)
        rr.bold = True
        rr.font.size = Pt(10)

    for row_data in table_rows:
        row = state_tbl.add_row()
        row.height = Mm(6.3)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        values = [
            row_data["nom_prenom"],
            row_data["qualite"],
            fmt_days(row_data["days"]),
            fmt_amount_fr(row_data["price"]) if row_data["price"] is not None else "",
            fmt_amount_fr(row_data["brut"]),
            fmt_amount_fr(row_data["prelev"]),
            fmt_amount_fr(row_data["versement"]),
        ]
        aligns = [
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ]
        for idx, value in enumerate(values):
            c = row.cells[idx]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = aligns[idx]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            rv = p.add_run(str(value))
            rv.font.size = Pt(10)

    total_row = state_tbl.add_row()
    total_row.height = Mm(7.0)
    total_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    totals_values = [
        "TOTAUX",
        "",
        fmt_days(total_days),
        "",
        fmt_amount_fr(total_brut),
        fmt_amount_fr(total_prelev),
        fmt_amount_fr(total_versement),
    ]
    for idx, value in enumerate(totals_values):
        c = total_row.cells[idx]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1, 2) else WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(value)
        rr.bold = True
        rr.font.size = Pt(10)

    p_words = doc.add_paragraph()
    p_words.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_words.paragraph_format.space_before = Pt(8)
    p_words.paragraph_format.space_after = Pt(10)
    p_words.add_run("Le présent état est arrêté à la somme de : ").font.size = Pt(11)
    p_words.add_run(amount_words_upper(total_versement)).font.size = Pt(11)

    p_date = doc.add_paragraph(f"{city} le : ..............................")
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after = Pt(0)
    for run in p_date.runs:
        run.font.size = Pt(11)

    # Start page 2 (exactly 2 pages)
    doc.add_page_break()

    # ------------------------------
    # Page 2: Justificatif de Versement
    # ------------------------------
    top_tbl = add_table_with_widths(doc, cols=2, col_widths_mm=[usable_w_mm * 0.62, usable_w_mm * 0.38])
    top_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(top_tbl, [usable_w_mm * 0.62, usable_w_mm * 0.38])
    remove_docx_table_borders(top_tbl)
    set_table_cell_margins(top_tbl, top_mm=0.0, bottom_mm=0.0, left_mm=0.0, right_mm=0.0)

    c_left = top_tbl.rows[0].cells[0]
    c_left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    clear_cell(c_left)
    p_logo = c_left.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)
    r_logo = p_logo.add_run("RCAR")
    r_logo.bold = True
    r_logo.font.size = Pt(24)

    if arabic_header:
        p_ar = c_left.add_paragraph(arabic_header)
        p_ar.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_ar.paragraph_format.space_before = Pt(0)
        p_ar.paragraph_format.space_after = Pt(0)
        for run in p_ar.runs:
            run.font.size = Pt(8)

    p_sub = c_left.add_paragraph("Regime Collectif d'Allocation de Retraite")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(0)
    for run in p_sub.runs:
        run.font.size = Pt(9)

    clear_cell(top_tbl.rows[0].cells[1])
    spacer_top = doc.add_paragraph()
    spacer_top.paragraph_format.space_before = Pt(0)
    spacer_top.paragraph_format.space_after = Pt(4)

    title_width = 95.0
    title_tbl = add_table_with_widths(doc, cols=1, col_widths_mm=[title_width])
    title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(title_tbl, [title_width])
    set_table_borders(title_tbl, sz=4)
    set_table_cell_margins(title_tbl, top_mm=0.8, bottom_mm=0.8, left_mm=1.0, right_mm=1.0)
    tcell = title_tbl.rows[0].cells[0]
    tcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_title = tcell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    r_title = p_title.add_run("Justificatif de Versement")
    r_title.bold = True
    r_title.font.size = Pt(14)

    spacer_after_title = doc.add_paragraph()
    spacer_after_title.paragraph_format.space_before = Pt(0)
    spacer_after_title.paragraph_format.space_after = Pt(4)

    info_tbl = add_table_with_widths(doc, cols=2, col_widths_mm=[usable_w_mm * 0.34, usable_w_mm * 0.66])
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(info_tbl, [usable_w_mm * 0.34, usable_w_mm * 0.66])
    remove_docx_table_borders(info_tbl)
    set_table_cell_margins(info_tbl, top_mm=0.0, bottom_mm=0.0, left_mm=0.0, right_mm=0.0)

    info_left = info_tbl.rows[0].cells[0]
    info_right = info_tbl.rows[0].cells[1]
    info_left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    info_right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    left_box_width = (usable_w_mm * 0.34) - 2.0
    left_box = info_left.add_table(rows=1, cols=1)
    left_box.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_fixed_layout(left_box, [left_box_width])
    set_table_borders(left_box, sz=4)
    set_table_cell_margins(left_box, top_mm=1.0, bottom_mm=1.0, left_mm=1.5, right_mm=1.5)
    lcell = left_box.rows[0].cells[0]
    lcell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    pl0 = lcell.paragraphs[0]
    pl0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pl0.paragraph_format.space_before = Pt(0)
    pl0.paragraph_format.space_after = Pt(8)
    rl0 = pl0.add_run("Dénomination")
    rl0.bold = True
    rl0.font.size = Pt(10)
    pl1 = lcell.add_paragraph(f"COMMUNE {commune}")
    pl1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl1.paragraph_format.space_before = Pt(5)
    pl1.paragraph_format.space_after = Pt(2)
    for run in pl1.runs:
        run.font.size = Pt(10)
    pl2 = lcell.add_paragraph(f"PROVINCE DE {province}")
    pl2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl2.paragraph_format.space_before = Pt(0)
    pl2.paragraph_format.space_after = Pt(0)
    for run in pl2.runs:
        run.font.size = Pt(10)

    fields_tbl = info_right.add_table(rows=3, cols=2)
    fields_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_fixed_layout(fields_tbl, [45.0, (usable_w_mm * 0.66) - 45.0])
    remove_docx_table_borders(fields_tbl)
    set_table_cell_margins(fields_tbl, top_mm=0.2, bottom_mm=0.2, left_mm=0.0, right_mm=0.0)

    for row in fields_tbl.rows:
        for c in row.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_a = fields_tbl.rows[0].cells[0].paragraphs[0]
    p_a.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_a.paragraph_format.space_before = Pt(0)
    p_a.paragraph_format.space_after = Pt(0)
    p_a.add_run("Numéro d'adhésion :").font.size = Pt(10)
    render_digit_boxes(fields_tbl.rows[0].cells[1], adhesion_number, boxes=8, box_width_mm=4.5, font_size_pt=10)

    p_y = fields_tbl.rows[1].cells[0].paragraphs[0]
    p_y.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_y.paragraph_format.space_before = Pt(0)
    p_y.paragraph_format.space_after = Pt(0)
    p_y.add_run("Année :").font.size = Pt(10)
    render_digit_boxes(fields_tbl.rows[1].cells[1], f"{year:04d}" if year > 0 else "", boxes=4, box_width_mm=4.5, font_size_pt=10)

    p_q = fields_tbl.rows[2].cells[0].paragraphs[0]
    p_q.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_q.paragraph_format.space_before = Pt(0)
    p_q.paragraph_format.space_after = Pt(0)
    p_q.add_run("Trimestre :").font.size = Pt(10)

    tri_cell = fields_tbl.rows[2].cells[1]
    tri_tbl = tri_cell.add_table(rows=1, cols=3)
    tri_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_fixed_layout(tri_tbl, [8.5, 10.0, 58.0])
    remove_docx_table_borders(tri_tbl)
    set_table_cell_margins(tri_tbl, top_mm=0.0, bottom_mm=0.0, left_mm=0.0, right_mm=0.0)
    render_digit_boxes(
        tri_tbl.rows[0].cells[0],
        str(quarter) if quarter in (1, 2, 3, 4) else "",
        boxes=1,
        box_width_mm=4.5,
        font_size_pt=10,
    )
    p_mois_label = tri_tbl.rows[0].cells[1].paragraphs[0]
    p_mois_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mois_label.paragraph_format.space_before = Pt(0)
    p_mois_label.paragraph_format.space_after = Pt(0)
    p_mois_label.add_run("MOIS").font.size = Pt(9)
    p_mois_val = tri_tbl.rows[0].cells[2].paragraphs[0]
    p_mois_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_mois_val.paragraph_format.space_before = Pt(0)
    p_mois_val.paragraph_format.space_after = Pt(0)
    rv_mois = p_mois_val.add_run(quarter_months_label(period_months))
    rv_mois.bold = True
    rv_mois.font.size = Pt(10)

    spacer_before_amounts = doc.add_paragraph()
    spacer_before_amounts.paragraph_format.space_before = Pt(0)
    spacer_before_amounts.paragraph_format.space_after = Pt(4)

    rg_rows = {
        "cotisation_salariale": None,
        "contribution_patronale": None,
        "cotisation_validation": None,
        "sous_total": None,
        "contribution_validation": None,
        "majoration_retard": None,
        "total_abc": None,
        "total_general": None,
    }
    rg_rows[justificatif_rg_key] = total_prelev
    rg_rows["sous_total"] = total_prelev
    rg_rows["total_abc"] = total_prelev
    rg_rows["total_general"] = total_prelev

    form_col_w = [usable_w_mm * 0.34, usable_w_mm * 0.33, usable_w_mm * 0.33]
    form_tbl = add_table_with_widths(doc, cols=3, col_widths_mm=form_col_w)
    form_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(form_tbl, form_col_w)
    set_table_borders(form_tbl, sz=4)
    set_table_cell_margins(form_tbl, top_mm=0.35, bottom_mm=0.35, left_mm=0.8, right_mm=0.8)

    h1 = form_tbl.rows[0]
    h1.height = Mm(7.5)
    h1.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    h1.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = h1.cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(0)
    rr0 = p0.add_run("Nature du Versement")
    rr0.bold = True
    rr0.font.size = Pt(10)
    h1.cells[1].merge(h1.cells[2])
    p01 = h1.cells[1].paragraphs[0]
    p01.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p01.paragraph_format.space_before = Pt(0)
    p01.paragraph_format.space_after = Pt(0)
    rr01 = p01.add_run("Montants en DH")
    rr01.bold = True
    rr01.font.size = Pt(10)

    h2 = form_tbl.add_row()
    h2.height = Mm(7.5)
    h2.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    h2.cells[0].text = ""
    p_rg = h2.cells[1].paragraphs[0]
    p_rg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rg.paragraph_format.space_before = Pt(0)
    p_rg.paragraph_format.space_after = Pt(0)
    rr_rg = p_rg.add_run("Régime Général (RG)")
    rr_rg.bold = True
    rr_rg.font.size = Pt(9)
    p_rc = h2.cells[2].paragraphs[0]
    p_rc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rc.paragraph_format.space_before = Pt(0)
    p_rc.paragraph_format.space_after = Pt(0)
    rr_rc = p_rc.add_run("Régime Complémentaire (RC)")
    rr_rc.bold = True
    rr_rc.font.size = Pt(9)

    form_rows = [
        ("Cotisation Salariale", "cotisation_salariale"),
        ("Contribution Patronale", "contribution_patronale"),
        ("Cotisation Validation", "cotisation_validation"),
        ("Sous-Total", "sous_total"),
        ("Contribution Validation", "contribution_validation"),
        ("Majoration de Retard", "majoration_retard"),
        ("Total (A + B + C)", "total_abc"),
        ("Total Général (RG + RC)", "total_general"),
    ]

    for label, key in form_rows:
        row = form_tbl.add_row()
        row.height = Mm(10.2)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for c in row.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pl = row.cells[0].paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pl.paragraph_format.space_before = Pt(0)
        pl.paragraph_format.space_after = Pt(0)
        rl = pl.add_run(label)
        rl.font.size = Pt(10)

        render_amount_boxes(row.cells[1], rg_rows.get(key), total_width_mm=form_col_w[1] - 2.0, boxes=11)
        render_amount_boxes(row.cells[2], None, total_width_mm=form_col_w[2] - 2.0, boxes=11)

    spacer_bottom = doc.add_paragraph()
    spacer_bottom.paragraph_format.space_before = Pt(0)
    spacer_bottom.paragraph_format.space_after = Pt(5)

    bottom_tbl = add_table_with_widths(doc, cols=2, col_widths_mm=[usable_w_mm * 0.5, usable_w_mm * 0.5])
    bottom_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(bottom_tbl, [usable_w_mm * 0.5, usable_w_mm * 0.5])
    set_table_borders(bottom_tbl, sz=4)
    set_table_cell_margins(bottom_tbl, top_mm=1.0, bottom_mm=1.0, left_mm=1.0, right_mm=1.0)

    b_row = bottom_tbl.rows[0]
    b_row.height = Mm(42)
    b_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for c in b_row.cells:
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    lrun = b_row.cells[0].paragraphs[0].add_run("Cadre réservé au RCAR")
    lrun.bold = True
    lrun.font.size = Pt(10)
    rrun = b_row.cells[1].paragraphs[0].add_run("Cachet et signature")
    rrun.bold = True
    rrun.font.size = Pt(10)

    doc.save(docx_path)


def generate_rcar_salariale_docx(payload: Dict[str, Any], docx_path: str) -> None:
    _generate_rcar_docx(
        payload,
        docx_path,
        expected_document_type="rcar-salariale",
        subtitle="COTISATION SALARIALE",
        prelevement_rate=0.06,
        justificatif_rg_key="cotisation_salariale",
    )


def generate_rcar_patronale_docx(payload: Dict[str, Any], docx_path: str) -> None:
    _generate_rcar_docx(
        payload,
        docx_path,
        expected_document_type="rcar-patronale",
        subtitle="COTISATION PATRONALE",
        prelevement_rate=0.12,
        justificatif_rg_key="contribution_patronale",
    )


def generate_generic_docx(payload: Dict[str, Any], docx_path: str) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm
    except ModuleNotFoundError as exc:
        raise RuntimeError("Missing Python dependency for Word generation. Install: python-docx") from exc

    document_type = str(payload.get("documentType") or "").strip()
    year = int(payload.get("year") or 0)
    month = payload.get("month")
    quarter = payload.get("quarter")
    report = payload.get("report") or {}
    report_rows = report.get("rows") or []
    options = payload.get("options") or {}
    ref = options or {}
    age_limit = int(options.get("rcarAgeLimit") or 60)

    if document_type == "recu-combined":
        generate_recu_docx(payload, docx_path)
        return
    if document_type == "demande-autorisation":
        generate_demande_autorisation_docx(payload, docx_path)
        return
    if document_type in ("certificat-paiement", "certificat-paiement-combined"):
        generate_certificat_paiement_docx(payload, docx_path)
        return
    if document_type == "ordre-paiement":
        generate_ordre_paiement_docx(payload, docx_path)
        return
    if document_type == "mandat-paiement":
        generate_mandat_paiement_docx(payload, docx_path)
        return
    if document_type == "bordereau":
        generate_bordereau_docx(payload, docx_path)
        return
    if document_type == "rcar-salariale":
        generate_rcar_salariale_docx(payload, docx_path)
        return
    if document_type == "rcar-patronale":
        generate_rcar_patronale_docx(payload, docx_path)
        return

    doc = Document()
    section0 = doc.sections[0]
    section0.page_width = Mm(210)
    section0.page_height = Mm(297)
    set_page_margins(section0, top_mm=12, bottom_mm=12, left_mm=12, right_mm=12)
    set_default_font(doc, font_name="Times New Roman", font_size_pt=10)

    title = DOC_TITLE_MAP.get(document_type, document_type.upper() if document_type else "DOCUMENT")
    add_centered_title(doc, title, font_size_pt=14)

    if year and month:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(f"Période: {str(month).zfill(2)}/{year}")
    elif year and quarter:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(f"Période: T{quarter}/{year}")

    ref_lines = []
    for label, key in [
        ("Chap", "chap"),
        ("Art", "art"),
        ("Prog", "prog"),
        ("Proj", "proj"),
        ("Ligne", "ligne"),
    ]:
        value = ref.get(key)
        if value:
            ref_lines.append(f"{label}: {value}")

    if ref_lines:
        doc.add_paragraph("Références budgétaires:")
        for line in ref_lines:
            doc.add_paragraph(line)

    if isinstance(report_rows, list) and report_rows:
        if year and month:
            _, end_date = month_start_end(year, int(month))
            ref_date = end_date
        else:
            ref_date = parse_date((report.get("period") or {}).get("quarterEndDate")) or date.today()

        rows_fin, totals = build_workers_financial_rows(report_rows, ref_date, age_limit)

        if rows_fin:
            doc.add_paragraph("Détail des ouvriers:")
            usable_w_mm = 210 - 24  # page width minus left/right margins (12mm each)
            col_widths_mm = [
                usable_w_mm * 0.30,
                usable_w_mm * 0.14,
                usable_w_mm * 0.07,
                usable_w_mm * 0.08,
                usable_w_mm * 0.12,
                usable_w_mm * 0.14,
                usable_w_mm * 0.15,
            ]
            table = add_table_with_widths(doc, cols=7, col_widths_mm=col_widths_mm)
            hdr = table.rows[0].cells
            headers = ["Nom et Prénom", "CIN", "Type", "Jours", "Brut", "Prélèvement", "Net"]
            for i, h in enumerate(headers):
                hdr[i].text = h

            for w in rows_fin:
                r = table.add_row().cells
                r[0].text = w["nom_prenom"]
                r[1].text = w["cin"]
                r[2].text = w["type"]
                r[3].text = str(w["days"])
                r[4].text = fmt_amount(w["gross"])
                r[5].text = fmt_amount(w["deduction"])
                r[6].text = fmt_amount(w["net"])

            doc.add_paragraph(
                f"Totaux - Jours: {int(totals['days'])} | Brut: {fmt_amount(totals['gross'])} | "
                f"Prélèvement: {fmt_amount(totals['deduction'])} | Net: {fmt_amount(totals['net'])}"
            )

    doc.save(docx_path)


def main() -> None:
    payload = parse_input_json()
    document_type = str(payload.get("documentType") or "").strip()
    output_dir = str(payload.get("outputDir") or "").strip()

    if not output_dir:
        raise ValueError("Missing required field: outputDir")
    if not document_type:
        raise ValueError("Missing required field: documentType")

    ensure_dir(output_dir)

    docx_name = build_docx_filename(document_type, payload)
    docx_path = os.path.join(output_dir, docx_name)
    generate_generic_docx({**payload, "documentType": document_type}, docx_path)
    result: Dict[str, Any] = {"success": True, "docxFileName": docx_name, "docxFilePath": docx_path}

    sys.stdout.write(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        error = {"success": False, "message": str(exc)}
        sys.stdout.write(json.dumps(error, ensure_ascii=False))
        sys.exit(1)
